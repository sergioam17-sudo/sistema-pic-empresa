#En la Versión 5.3 se incluye generación del documento en Word
#En la versión 5.4 se incluy las columnas de nombre del referente y el supervisor que aprueba
#En la vesión 5.5 Se incluye y mejora el análisis del dasboard
#En la versión 5.6 se incluye mejorar la visual de los valores del dasboard
# En la versión 5.7 se incluye mejor vsualización del dasboard del municipio
# En la versión 5.8 se incluye mejorar el informe de Word
# En la versión 5.9 se incluye el colocar en todas las tablas de trazabilidad y seguimiento las columnas de id_actividad y # nombre de la actividad
#En esta versión 6 se crear un certificado o acta de conformidad de las actividades realizadas por el referente la cual debe  # generarse de forma automática donde se indique para el municipio las actividades revisadas y aprobadas con un análisis   # de lo operativo y financiero revisado
# En la Versión 6.1 se incluye la columna de observaciones para el supervisor y e referente
# En la versión 6.2 se incluye el seguimiento administrativo del contrato que solo se hace solo por el supervisor del contrato
# En la versión 6.3 se incluye opción de rechazo por parte del referente al igual que el supervisor para que el municipio sudsane la novedad
# En la versión 6.4 se incluye tabla de secuencia para verificar los pagos realizados
# En la versión 6.5 se potencia el análisis de datos del tablero incluyendo la información de pagos con la tabla de secuencia
#En la versión 6.6 Se incluye en el informe Word del supervisión la información financiera de los pagos de los informes de supervisor

import streamlit as st
import pandas as pd
from streamlit_gsheets import GSheetsConnection



# URL corregida (sin parámetros de pestaña específicos)
URL_DB = "https://docs.google.com/spreadsheets/d/1jRdZX0gNfjWhlb86hHopVkJJrs_9bRIaulZDRzKR0pA/edit?usp=sharing"

# Crear conexión
conn = st.connection("gsheets", type=GSheetsConnection)

# --- FUNCIÓN PARA INICIALIZAR ENCABEZADOS (AUTOMÁTICO) ---

def init_excel_db():
    tablas = {
        "usuarios": ["id_usuario", "nombre_completo", "email", "password", "rol", "municipio_asignado"],
        "actividades_maestro": ["id_actividad", "nombre_actividad", "descripcion", "meta_global", "unidad_medida", "valor_total_actividad", "programa_responsable"],
        "subactividades": ["id_sub", "id_actividad", "nombre_subactividad", "valor_sub", "meta_sub", "unidad_medida_sub", "peso"],
        "asignacion_municipios": ["id_asig", "id_sub", "municipio", "num_contrato", "num_pagos", "valor_asignado", "meta_municipal", "unidad_medida_muni"],
        "seguimiento_pagos": ["id_seguimiento", "id_asig", "num_pago_actual", "avance_meta", "valor_calculado", "fecha_registro", "soporte_url", "estado", "referente_aprobador", "acta_referente", "observaciones_referente", "supervisor_aprobador", "motivo_rechazo", "chk_plan_trabajo", "chk_cronograma", "chk_personal", "chk_seg_social", "chk_inf_parcial", "chk_inf_final", "chk_polizas"],
        "secuencia": ["id_secuencia", "id_seguimiento", "id_asig", "municipio", "numero_contrato", "cp_nit_beneficiario", "numero_pagos", "primer_pago","ultimo_pago","total_pagado_oc","valor_cp","saldo_cp","porcentaje_ejecucion"]
    }
    
    for nombre, columnas in tablas.items():
        try:
            # TTL=0 para forzar lectura fresca
            df = conn.read(spreadsheet=URL_DB, worksheet=nombre, ttl=0)
            if df is None or df.empty:
                df_init = pd.DataFrame(columns=columnas)
                safe_update(nombre, df_init)
        except Exception:
            # Si la hoja no existe físicamente, la crea con los encabezados
            df_init = pd.DataFrame(columns=columnas)
            safe_update(nombre, df_init)


import time

# Nueva función para manejar el tráfico de 87 municipios al escribir
def safe_update(worksheet_name, df_final):
    max_retries = 3
    for i in range(max_retries):
        try:
            conn.update(spreadsheet=URL_DB, worksheet=worksheet_name, data=df_final)
            st.cache_data.clear() # Limpia la memoria local para que todos vean el cambio
            return True
        except Exception as e:
            if "429" in str(e):
                espera = (i + 1) * 5
                st.warning(f"⚠️ Servidor ocupado (Tráfico alto). Reintentando en {espera}s...")
                time.sleep(espera)
            else:
                st.error(f"Error al sincronizar: {e}")
                return False
    return False








# --- REEMPLAZO DE 'INSERT INTO' ---
def guardar_nuevo_usuario(nombre, email, clave, rol, muni):
    # 1. Leer datos actuales
    df_actual = conn.read(spreadsheet=URL_DB, worksheet="usuarios")
    
    # 2. Crear nueva fila (Pandas)
    nuevo_id = 1 if df_actual.empty else df_actual["id_usuario"].max() + 1
    nueva_fila = pd.DataFrame([[nuevo_id, nombre, email, clave, rol, muni]], 
                              columns=df_actual.columns)
    
    # 3. Unir y Subir
    df_final = pd.concat([df_actual, nueva_fila], ignore_index=True)
    
    success = safe_update("usuarios", df_final)
    if success:
        st.success("Usuario guardado y sincronizado.")









# --- CONFIGURACIÓN DE LECTURA ---
# --- CONFIGURACIÓN DE LECTURA OPTIMIZADA CON CACHÉ ---
# Se utiliza st.cache_data para que los 87 municipios consulten la RAM antes que el Drive
@st.cache_data(ttl=120)  # Mantiene los datos en memoria por 2 minutos para todos
def get_data_cached(nombre_hoja):
    return conn.read(spreadsheet=URL_DB, worksheet=nombre_hoja)

# --- CÓDIGO CORREGIDO Y OPTIMIZADO ---
def get_data(nombre_hoja, forzar=False):
    if forzar:
        # Si se acaba de guardar algo, limpiamos la caché y leemos directo
        st.cache_data.clear()
        try:
            return conn.read(spreadsheet=URL_DB, worksheet=nombre_hoja, ttl=0)
        except Exception:
            time.sleep(2)
            return conn.read(spreadsheet=URL_DB, worksheet=nombre_hoja, ttl=0)
    else:
        # Uso de la memoria local para consultas masivas
        try:
            return get_data_cached(nombre_hoja)
        except Exception as e:
            st.error(f"⚠️ Error de conexión en '{nombre_hoja}'. Reintentando lectura directa...")
            time.sleep(2)
            return conn.read(spreadsheet=URL_DB, worksheet=nombre_hoja, ttl=0)




# --- OPTIMIZACIÓN: Solo inicializar una vez por sesión para ahorrar cuota ---
if 'db_initialized' not in st.session_state:
    try:
        init_excel_db()
        st.session_state['db_initialized'] = True
    except Exception as e:
        if "429" in str(e):
            st.warning("⚠️ Google está procesando muchas solicitudes. Espera 30 segundos y refresca la página.")
        else:
            st.error(f"Error al conectar con la base de datos: {e}")



# --- LISTA DE MUNICIPIOS ---

municipios_santander = [
                "Aguada", "Albania", "Aratoca", "Barbosa", "Barichara", "Barrancabermeja", "Betulia", "Bolívar", 
                "Bucaramanga", "Cabrera", "California", "Capitanejo", "Carcasí", "Cepitá", "Cerrito", "Charalá", 
                "Charta", "Chima", "Chipatá", "Cimitarra", "Concepción", "Confines", "Contratación", "Coromoro", 
                "Curití", "El Carmen de Chucurí", "El Guacamayo", "El Peñón", "El Playón", "Encino", "Enciso", 
                "Floridablanca", "Florián", "Galán", "Gambita", "Girón", "Guaca", "Guadalupe", "Guapotá", "Guavatá", 
                "Güepsa", "Hato", "Jesús María", "Jordán", "La Belleza", "La Paz", "Landázuri", "Lebrija", "Los Santos", 
                "Macaravita", "Málaga", "Matanza", "Mogotes", "Molagavita", "Ocamonte", "Oiba", "Onzaga", "Palmar", 
                "Palmas del Socorro", "Páramo", "Piedecuesta", "Pinchote", "Puente Nacional", "Puerto Parra", 
                "Puerto Wilches", "Rionegro", "Sabana de Torres", "San Andrés", "San Benito", "San Gil", 
                "San Joaquín", "San José de Miranda", "San Miguel", "San Vicente de Chucurí", "Santa Bárbara", 
                "Santa Helena del Opón", "Simacota", "Socorro", "Suaita", "Sucre", "Suratá", "Tona", "Valle de San José", 
                "Vélez", "Vetas", "Villanueva", "Zapatoca"
            ]

# --- 2. CONFIGURACIÓN DE LA INTERFAZ ---
st.set_page_config(page_title="Sistema PIC - Unidad de Medida", layout="wide")

if 'user' not in st.session_state:
    st.sidebar.title("🔐 Acceso PIC Santander")
    user_input = st.sidebar.text_input("Email / Usuario")
    pass_input = st.sidebar.text_input("Contraseña", type="password")
    
    if st.sidebar.button("Ingresar"):
        
        try:
            # 1. Traemos la hoja de usuarios desde Google Sheets
            df_usuarios = get_data("usuarios")
            
            # 2. Filtramos el DataFrame buscando coincidencia
            user_match = df_usuarios[
                (df_usuarios['email'] == user_input) & 
                (df_usuarios['password'].astype(str) == str(pass_input))
            ]
             
            if not user_match.empty:
                # 3. Guardamos los datos en la sesión si hay coincidencia
                st.session_state['user'] = user_match.iloc[0]['email']
                st.session_state['rol'] = user_match.iloc[0]['rol']
                st.session_state['muni_asignado'] = user_match.iloc[0]['municipio_asignado']
                st.success(f"Bienvenido {user_match.iloc[0]['nombre_completo']}")
                st.rerun()
            else:
                st.sidebar.error("Usuario o contraseña incorrectos.")
        except Exception as e:
            st.sidebar.error(f"Error al conectar con la base de datos: {e}")



else:
    rol = st.session_state['rol']
    st.sidebar.info(f"**Usuario:** {st.session_state['user']}\n\n**Rol:** {rol}")
    
    # 🚀 CONTROL GLOBAL DE DATOS: Carga unificada y adicion de secuencia de pagos
    df_act_raw = get_data("actividades_maestro")
    df_sub_raw = get_data("subactividades")
    df_asig_raw = get_data("asignacion_municipios")
    df_pagos_raw = get_data("seguimiento_pagos")
    df_secuencia_raw = get_data("secuencia")  # <-- Nueva consulta indexada en memoria
    
    opciones = ["🏠 Dashboard", "📝 Ejecución"]


    if rol == "DEPARTAMENTO_PARAMETRIZADOR":
        opciones += ["⚙️ Parametrización", "⚖️ Revisión", "👤 Gestión Usuarios"]
    elif rol == "REFERENTE_DEPARTAMENTAL":
        opciones += ["⚖️ Revisión"]
    elif rol == "SUPERVISOR":
        opciones += ["⚖️ Revisión"] # O las opciones que definas para el supervisor

    menu = st.sidebar.radio("Navegación", opciones)


    # --- BOTÓN PARA CAMBIAR DE PERFIL (CERRAR SESIÓN) --- 
    st.sidebar.write("---")
    if st.sidebar.button("🔒 Cerrar Sesión / Cambiar Perfil"):
        # Limpia la sesión para volver al formulario de acceso 
        st.session_state.clear()
        st.rerun()


# --- MÓDULO: DASHBOARD DINÁMICO REINGENIERÍA ANALÍTICA ---
    if menu == "🏠 Dashboard":
        st.title(f"📊 Seguimiento - PIC Santander")
        st.caption(f"Sistema Integrado de Analítica para el Plan de Intervenciones Colectivas")

        
        # Validación estructural de seguridad para evitar caídas de la app
        if df_asig_raw.empty or df_sub_raw.empty:


            st.info("ℹ️ Esperando la parametrización de actividades y asignaciones municipales para proyectar métricas.")
        else:
            # Unión estructural: Asignación Municipio + Subactividad base
            df_master = df_asig_raw.merge(df_sub_raw, on="id_sub", suffixes=('_muni', '_sub'))
            
            # Cruzar con Actividades Maestro si existen datos para traer el Programa Responsable
            if not df_act_raw.empty:
                df_master = df_master.merge(df_act_raw[['id_actividad', 'nombre_actividad', 'programa_responsable']], on="id_actividad", how="left")
            
            # Tratamiento y consolidación de la tabla de ejecución (seguimiento_pagos)
            if not df_pagos_raw.empty:
                # Filtrar únicamente los pagos validados por el área técnica o aprobados finales
                df_pagos_validados = df_pagos_raw[df_pagos_raw['estado'].isin(['ACEPTADA', 'REVISADO_REFERENTE'])].copy()
                
                # Agrupación por id_asig para obtener la ejecución histórica acumulada
                df_pagos_agg = df_pagos_validados.groupby('id_asig').agg({
                    'valor_calculado': 'sum',
                    'avance_meta': 'sum'
                }).reset_index()
                df_pagos_agg.columns = ['id_asig', 'total_ejecutado_financiero', 'total_ejecutado_fisico']
                
                # Unir el acumulado al DataFrame Maestro
                df_master = df_master.merge(df_pagos_agg, on="id_asig", how="left")
            else:
                df_master['total_ejecutado_financiero'] = 0.0
                df_master['total_ejecutado_fisico'] = 0.0

            # Limpieza de valores nulos provocados por el merge de las actividades sin pagos reportados
            df_master['total_ejecutado_financiero'] = df_master['total_ejecutado_financiero'].fillna(0.0)
            df_master['total_ejecutado_fisico'] = df_master['total_ejecutado_fisico'].fillna(0.0)
            
            # Forzar casting numérico correcto de las variables del esquema original
            df_master['valor_asignado'] = df_master['valor_asignado'].astype(float)
            df_master['meta_municipal'] = df_master['meta_municipal'].astype(float)

            # -------------------------------------------------------------
            # 2. SEGMENTACIÓN LÓGICA DE INTERFAZ BASADA EN ROLES VIGENTES
            # -------------------------------------------------------------
            
            # =============================================================
            # PERFIL: MUNICIPIO_EJECUTOR (Vista de Autocontrol Técnico)
            # =============================================================
            if rol == "MUNICIPIO_EJECUTOR":
                muni_user = st.session_state.get('muni_asignado')
                st.subheader(f"📍 Panel de Fiscalización Local - Municipio: {muni_user}")
                
                # Filtrar la matriz para dejar exclusivamente las metas asignadas a la entidad actual
                df_muni_data = df_master[df_master['municipio'] == muni_user].copy()
                
                if df_muni_data.empty:
                    st.warning(f"⚠️ El municipio {muni_user} no cuenta con asignaciones presupuestales registradas en el sistema.")
                else:
                    # Métricas Financieras y Operativas Clave
                    muni_total_asig = df_muni_data['valor_asignado'].sum()
                    muni_total_ejec = df_muni_data['total_ejecutado_financiero'].sum()
                    muni_saldo_pendiente = muni_total_asig - muni_total_ejec
                    muni_porc_global = (muni_total_ejec / muni_total_asig * 100) if muni_total_asig > 0 else 0
                    
                    # Sistema de Semáforo basado en Percentiles Teóricos de Avance Contractual
                    if muni_porc_global < 33.3:
                        color_semaforo = "🔴 EJECUCIÓN CRÍTICA"
                        help_msg = "Alerta: El avance financiero se encuentra en el percentil inferior. Requiere intervención inmediata."
                    elif 33.3 <= muni_porc_global < 66.6:
                        color_semaforo = "🟡 EN ALERTA TEMPRANA"
                        help_msg = "Atención: Progreso intermedio. Monitorear rezagos en la carga de actas de soporte."
                    else:
                        color_semaforo = "🟢 CUMPLIMIENTO ÓPTIMO"
                        help_msg = "Ritmo de ejecución alineado con las metas del departamento."

                    # Renderizado Flexible de KPIs en Columnas Estables
                    kpi_m1, kpi_m2, kpi_m3, kpi_m4 = st.columns(4)
                    with kpi_m1:
                        st.caption("📋 Presupuesto Asignado")
                        st.markdown(f"### ${muni_total_asig:,.0f}")
                    with kpi_m2:
                        st.caption("✅ Total Ejecutado (OK)")
                        st.markdown(f"### ${muni_total_ejec:,.0f}")
                        st.caption(f"📈 {muni_porc_global:.1f}% del total")
                    with kpi_m3:
                        st.caption("📥 Saldo Líquido")
                        st.markdown(f"### ${muni_saldo_pendiente:,.0f}")
                    with kpi_m4:
                        st.caption("🚨 Estatus del Municipio")
                        st.markdown(f"#### {color_semaforo}")

                    
                    st.markdown("---")
                    
                    col_m_izq, col_m_der = st.columns(2)
                    
                    with col_m_izq:
                        st.write("#### 🚦 Semáforo de Eficiencia Presupuestal por Subactividad")
                        # Dataframe enfocado en el porcentaje de avance individual
                        df_semaforo_local = df_muni_data[['nombre_subactividad', 'valor_asignado', 'total_ejecutado_financiero']].copy()
                        df_semaforo_local['% Avance'] = (df_semaforo_local['total_ejecutado_financiero'] / df_semaforo_local['valor_asignado'].replace(0,1)) * 100
                        
                        df_semaforo_local['Estado'] = df_semaforo_local['% Avance'].apply(
                            lambda x: '🔴 Rezago Crítico' if x < 33.3 else ('🟡 En Alerta' if x < 66.6 else '🟢 Al día')
                        )
                        df_semaforo_local.columns = ['Subactividad', 'Asignado', 'Ejecutado', '% Eficiencia', 'Semáforo']
                        
                        # Formatear la visualización para evitar que se corten los textos con puntos suspensivos
                        st.dataframe(
                            df_semaforo_local.sort_values(by='% Eficiencia'),
                            use_container_width=True,
                            hide_index=True,
                            column_config={
                                "Subactividad": st.column_config.TextColumn(
                                    "Subactividad (Evidencia Técnica)",
                                    width="large",
                                    disabled=True
                                ),
                                "Asignado": st.column_config.NumberColumn("Asignado", format="$%,.0f"),
                                "Ejecutado": st.column_config.NumberColumn("Ejecutado", format="$%,.0f"),
                                "% Eficiencia": st.column_config.NumberColumn("% Efic.", format="%.1f%%")
                            }
                        )

                        
                    with col_m_der:
                        st.write("#### 🎯 Alerta de Rezago: Actividades con Mayor Brecha Presupuestal")
                        df_muni_data['brecha_financiera'] = df_muni_data['valor_asignado'] - df_muni_data['total_ejecutado_financiero']
                        df_rezagadas = df_muni_data.sort_values(by='brecha_financiera', ascending=False).head(5)
                        
                        if df_rezagadas['brecha_financiera'].sum() > 0:


                            import plotly.express as px
                            
                            # Inteligencia de Datos: Creamos una etiqueta corta para el eje visible sin dañar el texto original
                            df_rezagadas['subactividad_corta'] = df_rezagadas['nombre_subactividad'].apply(
                                lambda x: x[:37] + "..." if len(str(x)) > 40 else x
                            )
                            
                            # Construcción del Gráfico Avanzado de Plotly
                            fig_rezago_muni = px.bar(
                                df_rezagadas,
                                x='brecha_financiera',
                                y='subactividad_corta',
                                orientation='h',
                                title="Top Subactividades con Mayor Brecha Financiera",
                                labels={'brecha_financiera': 'Brecha Pendiente ($)', 'subactividad_corta': 'Subactividad'},
                                # CRÍTICO: El hover mostrará la variable 'nombre_subactividad' ORIGINAL COMPLETA
                                hover_data={
                                    'subactividad_corta': False, # Ocultamos la versión recortada del tooltip
                                    'nombre_subactividad': True,  # Mostramos el nombre descriptivo completo
                                    'brecha_financiera': ':$,.2f' # Formato monetario profesional con decimales
                                }
                            )
                            
                            # Optimización del Layout para asegurar simetría visual y espacio suficiente
                            fig_rezago_muni.update_layout(
                                yaxis={
                                    'categoryorder': 'total ascending',
                                    'tickmode': 'linear'
                                },
                                margin=dict(l=220, r=30, t=50, b=40),
                                height=380,
                                hoverlabel=dict(
                                    bgcolor="white",
                                    font_size=12,
                                    font_family="Arial"
                                )
                            )
                            st.plotly_chart(fig_rezago_muni, use_container_width=True)


                        else:
                            st.success("🎉 ¡Excelente! El municipio ha ejecutado el 100% de los recursos disponibles.")


                    # Comparación del Porcentaje Físico vs Financiero con Ingeniería de Visualización
                    st.write("#### 📈 Balance de Ejecución: Financiero vs. Avance de Metas Físicas")
                    df_muni_data['% Eficiencia Financiera'] = (df_muni_data['total_ejecutado_financiero'] / df_muni_data['valor_asignado'].replace(0,1)) * 100
                    df_muni_data['% Avance Metas Físicas'] = (df_muni_data['total_ejecutado_fisico'] / df_muni_data['meta_municipal'].replace(0,1)) * 100
                    
                    df_balance_muni = df_muni_data[['nombre_subactividad', '% Eficiencia Financiera', '% Avance Metas Físicas']].copy()
                    
                    import plotly.express as px
                    
                    # 1. Crear etiquetas cortas para el eje X visible y evitar colisiones tipográficas
                    df_balance_muni['subactividad_corta'] = df_balance_muni['nombre_subactividad'].apply(
                        lambda x: str(x)[:22] + "..." if len(str(x)) > 25 else str(x)
                    )
                    
                    # 2. Formateador HTML para segmentar el texto del Tooltip flotante en párrafos limpios
                    def formatear_linea_hover(texto):
                        texto_str = str(texto)
                        chunks = [texto_str[i:i+50] for i in range(0, len(texto_str), 50)]
                        return "<br>".join(chunks)
                        
                    df_balance_muni['hover_completo'] = df_balance_muni['nombre_subactividad'].apply(formatear_linea_hover)
                    
                    # 3. Pivotar/Melt del DataFrame para que Plotly pinte las dos líneas correctamente con su leyenda
                    df_melted = df_balance_muni.melt(
                        id_vars=['subactividad_corta', 'hover_completo'], 
                        value_vars=['% Eficiencia Financiera', '% Avance Metas Físicas'],
                        var_name='Dimensión Evaluada', 
                        value_name='Porcentaje'
                    )
                    
                    # 4. Construcción interactiva del gráfico de líneas
                    fig_line_muni = px.line(
                        df_melted,
                        x='subactividad_corta',
                        y='Porcentaje',
                        color='Dimensión Evaluada',
                        markers=True,
                        title="<b>Análisis de Convergencia: Recursos Devengados vs. Metas Físicas (%)</b>",
                        labels={'subactividad_corta': 'Subactividades del PIC', 'Porcentaje': 'Porcentaje Realizado (%)'},
                        hover_data={
                            'subactividad_corta': False,      # Ocultamos el nombre corto recortado
                            'hover_completo': True,           # Desplegamos el nombre original estructurado
                            'Porcentaje': ':.1f%'             # Formato con un solo decimal de precisión
                        }
                    )
                    
                    # 5. Ajustes estéticos y de inclinación de etiquetas para máxima legibilidad
                    fig_line_muni.update_layout(
                        xaxis=dict(tickangle=-35, tickmode='linear'),
                        margin=dict(l=50, r=30, t=60, b=100),
                        height=420,
                        title_font=dict(size=14, family="Arial"),
                        legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1),
                        hoverlabel=dict(
                            bgcolor="#F8F9FA",
                            font_size=12,
                            font_family="Arial"
                        )
                    )
                    st.plotly_chart(fig_line_muni, use_container_width=True)


            # =============================================================
            # PERFILES: PARAMETRIZADOR, REFERENTE Y SUPERVISOR (ALTA GERENCIA)
            # =============================================================
            else:
                # Creación de Pestañas Macro solicitadas para separar el análisis del Departamento y del Municipio
                tab_departamento, tab_municipios_analitica = st.tabs([
                    "🏢 Macrodashboard: Análisis Departamental", 
                    "📍 Analítica Desagregada por Municipio"
                ])
                
                # --- CONTROL DE FILTRO EN LA BARRA LATERAL (USANDO VARIABLES EXISTENTES) ---
                st.sidebar.markdown("### 🎛️ Filtros Estructurales")
                
                # Extraer dinámicamente los números de cuota/pagos existentes en la base de datos de pagos
                if not df_pagos_raw.empty:
                    lista_pagos = ["TODOS"] + sorted(df_pagos_raw['num_pago_actual'].dropna().unique().astype(int).tolist())
                else:
                    lista_pagos = ["TODOS", 1, 2, 3]
                
                filtro_fase_pago = st.sidebar.selectbox("Fase de Seguimiento (Número de Pago)", lista_pagos)
                
                # Aplicar filtrado analítico por el número de cuota seleccionada si aplica
                if filtro_fase_pago != "TODOS" and not df_pagos_raw.empty:
                    df_pagos_filtrados = df_pagos_raw[
                        (df_pagos_raw['estado'].isin(['ACEPTADA', 'REVISADO_REFERENTE'])) & 
                        (df_pagos_raw['num_pago_actual'].astype(int) == int(filtro_fase_pago))
                    ].copy()
                    
                    df_p_agg_f = df_pagos_filtrados.groupby('id_asig').agg({
                        'valor_calculado': 'sum', 'avance_meta': 'sum'
                    }).reset_index()
                    df_p_agg_f.columns = ['id_asig', 'total_ejecutado_financiero', 'total_ejecutado_fisico']
                    
                    # Re-ensamblar la sábana maestra filtrada por la fase del desembolso
                    df_master_f = df_asig_raw.merge(df_sub_raw, on="id_sub", suffixes=('_muni', '_sub'))
                    if not df_act_raw.empty:
                        df_master_f = df_master_f.merge(df_act_raw[['id_actividad', 'nombre_actividad', 'programa_responsable']], on="id_actividad", how="left")
                    df_master_f = df_master_f.merge(df_p_agg_f, on="id_asig", how="left")
                    df_master_f['total_ejecutado_financiero'] = df_master_f['total_ejecutado_financiero'].fillna(0.0)
                    df_master_f['total_ejecutado_fisico'] = df_master_f['total_ejecutado_fisico'].fillna(0.0)
                else:
                    df_master_f = df_master.copy()

                # Forzar recálculo de porcentajes sobre la matriz final filtrada
                df_master_f['valor_asignado'] = df_master_f['valor_asignado'].astype(float)
                df_master_f['total_ejecutado_financiero'] = df_master_f['total_ejecutado_financiero'].astype(float)

                # =============================================================
                # PESTAÑA 1: CÁLCULOS MACRO DEPARTAMENTALES
                # =============================================================
                with tab_departamento:
                    st.subheader("🌐 Visión de Control Fiscal y Operativo de Santander")
                    
                    dep_total_pic = df_act_raw['valor_total_actividad'].sum() if not df_act_raw.empty else 0
                    dep_total_asig = df_master_f['valor_asignado'].sum()
                    dep_total_ejec = df_master_f['total_ejecutado_financiero'].sum()
                    dep_saldo_bolsa = dep_total_pic - dep_total_asig
                    dep_eficiencia = (dep_total_ejec / dep_total_asig * 100) if dep_total_asig > 0 else 0
                    
                    # Estructura de columnas para evitar el colapso del contenedor nativo st.metric
                    c_dep1, c_dep2, c_dep3, c_dep4 = st.columns(4)
                    with c_dep1:
                        st.caption("📋 Techo Total PIC")
                        st.markdown(f"### ${dep_total_pic:,.0f}")
                    with c_dep2:
                        st.caption("📍 Total Asignado Municipios")
                        st.markdown(f"### ${dep_total_asig:,.0f}")
                    with c_dep3:
                        st.caption("💰 Ejecución Financiera Real")
                        st.markdown(f"### ${dep_total_ejec:,.0f}")
                        st.caption(f"📈 {dep_eficiencia:.1f}% Eficiencia")
                    with c_dep4:
                        st.caption("💼 Saldo PIC Sin Asignar")
                        st.markdown(f"### ${dep_saldo_bolsa:,.0f}")
                    
                    st.markdown("---")
                    
                    col_dep_izq, col_dep_der = st.columns(2)
                    
                    with col_dep_izq:
                        st.write("#### 📉 Los 10 Municipios Más Rezagados en la Ejecución")
                        # Agrupación por municipio para medir su rendimiento porcentual
                        df_muni_perf = df_master_f.groupby('municipio').agg({
                            'valor_asignado': 'sum',
                            'total_ejecutado_financiero': 'sum'
                        }).reset_index()
                        df_muni_perf['% Ejecución'] = (df_muni_perf['total_ejecutado_financiero'] / df_muni_perf['valor_asignado'].replace(0,1)) * 100
                        
                        # Ordenar de menor a mayor porcentaje para mostrar los municipios que van más "quedados"
                        df_muni_rezagados = df_muni_perf.sort_values(by='% Ejecución', ascending=True).head(10)
                        
                        import plotly.express as px
                        fig_dep_rezago = px.bar(
                            df_muni_rezagados,
                            x='municipio',
                            y='% Ejecución',
                            title="Top 10 Municipios Más Rezagados en la Ejecución",
                            labels={'municipio': 'Municipio', '% Ejecución': 'Porcentaje de Ejecución (%)'},
                            color_discrete_sequence=["#EF4444"],
                            text_auto='.1f',  # Muestra el valor exacto encima de cada barra
                            hover_data={'municipio': True, '% Ejecución': ':.2f%'}
                        )
                        fig_dep_rezago.update_layout(margin=dict(l=40, r=20, t=40, b=40))
                        st.plotly_chart(fig_dep_rezago, use_container_width=True)
                        
                        st.caption("⚠️ Alerta Gerencial: Municipios ordenados de menor a mayor porcentaje de avance presupuestal.")

                    with col_dep_der:
                        st.write("#### 📋 Matriz de Semáforos Territoriales")
                        df_muni_perf['Semáforo'] = df_muni_perf['% Ejecución'].apply(
                            lambda x: "🔴 Crítico (<33.3%)" if x < 33.3 else ("🟡 Alerta (33%-66%)" if x < 66.6 else "🟢 Óptimo")
                        )
                        df_grid_alertas = df_muni_perf[['municipio', 'valor_asignado', '% Ejecución', 'Semáforo']].copy()
                        df_grid_alertas.columns = ['Municipio', 'Total Asignado', '% Eficiencia', 'Estatus Contractual']
                        st.dataframe(df_grid_alertas.sort_values(by='% Eficiencia'), use_container_width=True, hide_index=True)

                    st.markdown("---")
                    st.write("#### 🩺 Ejecución Presupuestal por Programa de Salud Pública")
                    if 'programa_responsable' in df_master_f.columns:
                        df_prog_analysis = df_master_f.groupby('programa_responsable').agg({
                            'valor_asignado': 'sum',
                            'total_ejecutado_financiero': 'sum'
                        }).reset_index()
                        df_prog_analysis['% Avance'] = (df_prog_analysis['total_ejecutado_financiero'] / df_prog_analysis['valor_asignado'].replace(0,1)) * 100
                        st.dataframe(df_prog_analysis.sort_values(by='% Avance'), use_container_width=True, hide_index=True)

                # =============================================================
                # PESTAÑA 2: ANÁLISIS ESPECÍFICO DESAGREGADO POR MUNICIPIO
                # =============================================================
                with tab_municipios_analitica:
                    st.subheader("🔎 Filtro de Precisión: Auditoría y Desglose Territorial")
                    
                    lista_muni_selector = sorted(df_master_f['municipio'].unique().tolist())
                    sel_muni_analisis = st.selectbox("Seleccione el Municipio a Fiscalizar:", lista_muni_selector, key="sb_muni_analitica")
                    
                    df_muni_especifico = df_master_f[df_master_f['municipio'] == sel_muni_analisis].copy()
                    
                    if df_muni_especifico.empty:
                        st.info("No se registran datos indexados para el municipio seleccionado.")
                    else:
                        m_asig = df_muni_especifico['valor_asignado'].sum()
                        m_ejec = df_muni_especifico['total_ejecutado_financiero'].sum()
                        m_porc = (m_ejec / m_asig * 100) if m_asig > 0 else 0
                        
                        col_an1, col_an2, col_an3 = st.columns(3)
                        with col_an1:
                            st.caption("📌 Monto Asignado")
                            st.markdown(f"### ${m_asig:,.0f}")
                        with col_an2:
                            st.caption("💵 Monto Ejecutado Aceptado")
                            st.markdown(f"### ${m_ejec:,.0f}")
                        with col_an3:
                            st.caption("📊 Porcentaje de Eficiencia")
                            st.markdown(f"### {m_porc:.2f}%")
                        
                        st.markdown("#### Desglose Técnico de Metas en el Municipio")
                        
                        # Calcular brecha real en pesos para identificar la actividad más rezagada
                        df_muni_especifico['Brecha Financiera ($)'] = df_muni_especifico['valor_asignado'] - df_muni_especifico['total_ejecutado_financiero']
                        df_muni_especifico['% Avance'] = (df_muni_especifico['total_ejecutado_financiero'] / df_muni_especifico['valor_asignado'].replace(0,1)) * 100
                        

                        df_tabla_muni_det = df_muni_especifico[[
                            'id_actividad', 'nombre_actividad', 'nombre_subactividad', 
                            'meta_municipal', 'total_ejecutado_fisico', 
                            'valor_asignado', 'total_ejecutado_financiero', 'Brecha Financiera ($)', '% Avance'
                        ]].copy()
                        
                        df_tabla_muni_det.columns = [
                            'ID Actividad', 'Actividad Maestro', 'Subactividad', 'Meta Programada', 'Avance Físico', 
                            'Presupuesto Asignado', 'Presupuesto Ejecutado', 'Brecha Económica', '% Eficiencia'
                        ]
                        


                        st.dataframe(df_tabla_muni_det.sort_values(by='% Eficiencia'), use_container_width=True, hide_index=True)
                        
                        # Extraer e inyectar el nombre de la subactividad que registra el mayor rezago presupuestal en el territorio
                        actividad_mas_critica = df_muni_especifico.sort_values(by='Brecha Financiera ($)', ascending=False).iloc[0]['nombre_subactividad']
                        st.error(f"⚠️ **Estrategia en Salud Pública:** En el municipio de **{sel_muni_analisis}**, la subactividad con mayor rezago presupuestal es **{actividad_mas_critica}**. Se aconseja agilizar las mesas de auditoría técnica.")


            
                # ==============================================================================
                # INCLUSIÓN ANALÍTICA: TRACKING DE FRECUENCIA Y COMPORTAMIENTO DE PAGOS
                # ==============================================================================
                st.write("---")
                st.header("💰 Trazabilidad de Flujo de Caja: Secuencia de Pagos")
                st.caption("Evaluación de la frecuencia de desembolsos, absorción presupuestal y sostenibilidad del PIC territorial.")

                if df_secuencia_raw is None or df_secuencia_raw.empty:
                    st.info("ℹ️ No se registran secuencias de pagos indexadas en la base de datos para computar indicadores.")
                else:
                    # 1. Casting Defensivo y Normalización de Tipos de Datos (Data Science Standard)
                    df_sec = df_secuencia_raw.copy()
                    df_sec['valor_cp'] = pd.to_numeric(df_sec['valor_cp'], errors='coerce').fillna(0.0)
                    df_sec['total_pagado_oc'] = pd.to_numeric(df_sec['total_pagado_oc'], errors='coerce').fillna(0.0)
                    df_sec['saldo_cp'] = pd.to_numeric(df_sec['saldo_cp'], errors='coerce').fillna(0.0)
                    df_sec['porcentaje_ejecucion'] = pd.to_numeric(df_sec['porcentaje_ejecucion'], errors='coerce').fillna(0.0)
                    df_sec['numero_pagos'] = pd.to_numeric(df_sec['numero_pagos'], errors='coerce').fillna(0).astype(int)

                    # 2. Control de Ámbito y Filtros Adaptativos por Rol de Seguridad
                    if rol == "MUNICIPIO_EJECUTOR":
                        muni_user_sec = st.session_state.get('muni_asignado')
                        st.info(f"📍 Restricción perimetral activa: Visualizando datos exclusivos del Municipio de **{muni_user_sec}**.")
                        df_sec_filtrado = df_sec[df_sec['municipio'] == muni_user_sec].copy()
                        vista_general_activa = False
                    else:
                        # Perfiles Directivos (Parametrizador, Referente, Supervisor): Selector jerárquico General vs Local
                        lista_municipios_sec = ["TODOS (Consolidado Departamental)"] + sorted(df_sec['municipio'].dropna().unique().tolist())
                        seleccion_muni_sec = st.selectbox("🎛️ Filtrar Comportamiento de Pagos por Jurisdicción:", lista_municipios_sec, key="sb_secuencia_auditor")
                        
                        if seleccion_muni_sec == "TODOS (Consolidado Departamental)":
                            df_sec_filtrado = df_sec.copy()
                            vista_general_activa = True
                        else:
                            df_sec_filtrado = df_sec[df_sec['municipio'] == seleccion_muni_sec].copy()
                            vista_general_activa = False

                    # 3. Renderizado de Métricas de Absorción y Frecuencia
                    if df_sec_filtrado.empty:
                        st.warning("⚠️ No se localizaron registros de dispersión de giros bajo los parámetros del filtro seleccionado.")
                    else:
                        # Cálculos macro-estadísticos ponderados
                        total_contratos_sec = df_sec_filtrado['numero_contrato'].nunique()
                        total_giros_emitidos = df_sec_filtrado['numero_pagos'].sum()
                        total_valor_cp_sec = df_sec_filtrado['valor_cp'].sum()
                        total_pagado_oc_sec = df_sec_filtrado['total_pagado_oc'].sum()
                        total_saldo_cp_sec = df_sec_filtrado['saldo_cp'].sum()
                        
                        eficiencia_ponderada_sec = (total_pagado_oc_sec / total_valor_cp_sec * 100) if total_valor_cp_sec > 0 else 0.0
                        promedio_giros_contrato = df_sec_filtrado['numero_pagos'].mean()
                        ticket_promedio_giro_oc = df_sec_filtrado['total_pagado_oc'].mean()

                        # Paneles KPI Estables y Limpios
                        kpi_sec1, kpi_sec2, kpi_sec3, kpi_sec4 = st.columns(4)
                        with kpi_sec1:
                            st.caption("📋 Contratos Auditados")
                            st.markdown(f"### {total_contratos_sec:,}")
                            st.caption(f"Dispersión: {promedio_giros_contrato:.1f} pagos/contrato")
                        with kpi_sec2:
                            st.caption("🔢 Cuotas/Pagos Realizados")
                            st.markdown(f"### {total_giros_emitidos:,} Giros")
                            st.caption(f"Ticket Prom. Línea: ${ticket_promedio_giro_oc:,.0f}")
                        with kpi_sec3:
                            st.caption("💰 Capital Absorbido (OC)")
                            st.markdown(f"### ${total_pagado_oc_sec:,.0f}")
                            st.caption(f"Techo Comprometido (CP): ${total_valor_cp_sec:,.0f}")
                        with kpi_sec4:
                            st.caption("📈 Coeficiente de Absorción")
                            st.markdown(f"### {eficiencia_ponderada_sec:.2f}%")
                            st.caption(f"Reserva Líquida en Caja: ${total_saldo_cp_sec:,.0f}")

                        st.markdown("---")

                        # 4. Ingeniería de Visualización Interactiva (Plotly Express de Alta Densidad)
                        import plotly.express as px
                        col_sec_izq, col_sec_der = st.columns(2)

                        with col_sec_izq:
                            if vista_general_activa:
                                st.write("#### 📊 Frecuencia Acumulada de Pagos por Ente Territorial")
                                df_muni_sec_agg = df_sec_filtrado.groupby('municipio').agg({
                                    'numero_pagos': 'sum',
                                    'total_pagado_oc': 'sum',
                                    'valor_cp': 'sum'
                                }).reset_index().sort_values(by='numero_pagos', ascending=False).head(10)

                                fig_vol_giros = px.bar(
                                    df_muni_sec_agg,
                                    x='numero_pagos',
                                    y='municipio',
                                    orientation='h',
                                    title="Top 10 Municipios con Mayor Volumen de Transacciones",
                                    labels={'numero_pagos': 'Número de Pagos Realizados', 'municipio': 'Municipio'},
                                    color='total_pagado_oc',
                                    color_continuous_scale=px.colors.sequential.Bluyl,
                                    hover_data={'total_pagado_oc': ':$,.2f', 'valor_cp': ':$,.2f'}
                                )
                                fig_vol_giros.update_layout(yaxis={'categoryorder': 'total ascending'}, margin=dict(l=150, r=20, t=50, b=40), height=380)
                                st.plotly_chart(fig_vol_giros, use_container_width=True)
                            else:
                                st.write("#### 📊 Maduración de Giros por Línea Contractual Local")
                                df_sec_filtrado['contrato_tag'] = df_sec_filtrado['numero_contrato'].apply(lambda x: f"Contrato: {x}")
                                
                                fig_vol_giros = px.bar(
                                    df_sec_filtrado,
                                    x='numero_pagos',
                                    y='contrato_tag',
                                    orientation='h',
                                    title="Frecuencia de Pagos Reportados por Eje de Contratación",
                                    labels={'numero_pagos': 'Número de Pagos Registrados', 'contrato_tag': 'Estructura Contractual'},
                                    color='porcentaje_ejecucion',
                                    color_continuous_scale=px.colors.sequential.Viridis,
                                    hover_data={'total_pagado_oc': ':$,.2f', 'saldo_cp': ':$,.2f', 'porcentaje_ejecucion': ':.2f%'}
                                )
                                fig_vol_giros.update_layout(yaxis={'categoryorder': 'total ascending'}, margin=dict(l=150, r=20, t=50, b=40), height=380)
                                st.plotly_chart(fig_vol_giros, use_container_width=True)

                        with col_sec_der:
                            st.write("#### 📈 Balance Estructural de Presupuestos (Devengado vs. Reserva)")
                            if vista_general_activa:
                                df_bal_muni_sec = df_sec_filtrado.groupby('municipio').agg({
                                    'total_pagado_oc': 'sum',
                                    'saldo_cp': 'sum'
                                }).reset_index().sort_values(by='total_pagado_oc', ascending=False).head(8)

                                df_melted_sec = df_bal_muni_sec.melt(id_vars=['municipio'], value_vars=['total_pagado_oc', 'saldo_cp'],
                                                                     var_name='Estado Financiero', value_name='Presupuesto ($)')
                                df_melted_sec['Estado Financiero'] = df_melted_sec['Estado Financiero'].map({'total_pagado_oc': 'Total Girado', 'saldo_cp': 'Saldo en Reserva'})

                                fig_balance_sec = px.bar(
                                    df_melted_sec,
                                    x='municipio',
                                    y='Presupuesto ($)',
                                    color='Estado Financiero',
                                    title="Composición Presupuestal en Top 8 Municipios de Alta Ejecución",
                                    labels={'municipio': 'Entidad Territorial', 'Presupuesto ($)': 'Monto Económico ($)'},
                                    barmode='stack',
                                    color_discrete_sequence=["#059669", "#DC2626"]
                                )
                                fig_balance_sec.update_layout(xaxis=dict(tickangle=-25), margin=dict(l=50, r=20, t=50, b=80), height=380)
                                st.plotly_chart(fig_balance_sec, use_container_width=True)
                            else:
                                df_melted_sec = df_sec_filtrado.melt(id_vars=['numero_contrato'], value_vars=['total_pagado_oc', 'saldo_cp'],
                                                                     var_name='Estado Financiero', value_name='Presupuesto ($)')
                                df_melted_sec['Estado Financiero'] = df_melted_sec['Estado Financiero'].map({'total_pagado_oc': 'Total Girado', 'saldo_cp': 'Saldo en Reserva'})

                                fig_balance_sec = px.bar(
                                    df_melted_sec,
                                    x='numero_contrato',
                                    y='Presupuesto ($)',
                                    color='Estado Financiero',
                                    title="Balance Presupuestal Discriminado por Contrato Local",
                                    labels={'numero_contrato': 'Código Contractual', 'Presupuesto ($)': 'Valor ($)'},
                                    barmode='group',
                                    color_discrete_sequence=["#2563EB", "#D97706"]
                                )
                                fig_balance_sec.update_layout(margin=dict(l=50, r=20, t=50, b=40), height=380)
                                st.plotly_chart(fig_balance_sec, use_container_width=True)

                        # 5. Trazabilidad de Puntos de Control y Hitos Temporales
                        st.write("#### 📑 Matriz Cronológica de Hitos de Liquidación de Giros")
                        df_grid_sec = df_sec_filtrado[[
                            'municipio', 'numero_contrato', 'cp_nit_beneficiario', 'numero_pagos', 
                            'primer_pago', 'ultimo_pago', 'valor_cp', 'total_pagado_oc', 'porcentaje_ejecucion'
                        ]].copy()
                        
                        df_grid_sec.columns = [
                            'Municipio', 'N° Contrato', 'NIT/Cédula Beneficiario', 'Giros Totales', 
                            'Hito Primer Pago', 'Hito Último Pago', 'Valor Registro CP', 'Total Desembolsado', 'Eficiencia Real'
                        ]
                        
                        st.dataframe(
                            df_grid_sec.sort_values(by='Giros Totales', ascending=False),
                            use_container_width=True,
                            hide_index=True,
                            column_config={
                                "Valor Registro CP": st.column_config.NumberColumn("Valor Registro CP", format="$%,.2f"),
                                "Total Desembolsado": st.column_config.NumberColumn("Total Desembolsado", format="$%,.2f"),
                                "Eficiencia Real": st.column_config.NumberColumn("Eficiencia Real", format="%.2f%%")
                            }
                        )

    



    # --- MÓDULO: PARAMETRIZACIÓN ---
    if menu == "⚙️ Parametrización":
        st.title("⚙️ Módulo de Parametrización")
        tab1, tab2, tab3 = st.tabs(["1. Registro de Actividades", "2. Configuración de Subactividades", "3. Asignación a Municipios"])

        # TAB 1: ACTIVIDADES
        with tab1:
            st.subheader("Crear Nueva Actividad General")
            with st.form("form_actividad"):
                c1, c2 = st.columns(2)
                nombre_a = c1.text_input("Nombre de la Actividad")
                prog = c1.text_input("Programa Responsable")
                
                val_total = c2.number_input("Valor Total Actividad ($)", min_value=0.0)
                cc1, cc2 = c2.columns(2)
                meta_a = cc1.number_input("Meta Global", min_value=0.0)
                unidad_a = cc2.text_input("Unidad (Ej: Personas, Talleres)")
                
                desc_a = st.text_area("Descripción de la Actividad")
                
                if st.form_submit_button("💾 Guardar Actividad"):
                    df_act = get_data("actividades_maestro")
                    nuevo_id = 1 if df_act.empty else df_act['id_actividad'].max() + 1
                    
                    nueva_fila = pd.DataFrame([{
                        "id_actividad": nuevo_id,
                        "nombre_actividad": nombre_a, "descripcion": desc_a,
                        "meta_global": meta_a, "unidad_medida": unidad_a,
                        "valor_total_actividad": val_total, "programa_responsable": prog
                    }])
                    
                    df_final = pd.concat([df_act, nueva_fila], ignore_index=True)
                    safe_update("actividades_maestro", df_final)
                    st.success("✅ Actividad guardada en Sheets.")
                    st.rerun()



# --- TABLA DE ACTIVIDADES CON OPCIÓN DE ELIMINAR ---
            st.write("---")
            st.subheader("📋 Actividades Generales Registradas")
            
            df_maestro = get_data("actividades_maestro")
            
            if not df_maestro.empty:
                # Mostramos la tabla profesional
                st.dataframe(df_maestro[['id_actividad', 'nombre_actividad', 'programa_responsable', 'valor_total_actividad', 'meta_global', 'unidad_medida']], use_container_width=True)
                
                # Opción para eliminar por si hubo error
                with st.expander("🗑️ Zona de eliminación (Usar con precaución)"):
                    id_a_borrar = st.number_input("Ingrese el ID de la actividad a eliminar:", min_value=1, step=1)
                    if st.button("Eliminar Actividad Permanentemente"):
                        # Validamos en el DataFrame de subactividades
                        df_sub_check = get_data("subactividades")
                        check_sub = df_sub_check[df_sub_check['id_actividad'] == id_a_borrar]
                        
                        if not check_sub.empty:
                            st.error("⚠️ No se puede eliminar: Esta actividad tiene subactividades vinculadas.")
                        else:
                            df_maestro_del = get_data("actividades_maestro")
                            df_maestro_del = df_maestro_del[df_maestro_del['id_actividad'] != id_a_borrar]
                            safe_update("actividades_maestro", df_maestro_del)
                            st.warning(f"Actividad {id_a_borrar} eliminada del Excel.")
                            st.rerun()
	
            else:
                st.info("No hay actividades registradas todavía.")




        # TAB 2: SUBACTIVIDADES
        with tab2:
            st.subheader("Desglose de Subactividades")
            df_act = get_data("actividades_maestro")

            if not df_act.empty:
                nombres_act = {row['id_actividad']: row['nombre_actividad'] for _, row in df_act.iterrows()}
                act_id = st.selectbox("Seleccione Actividad Padre:", df_act['id_actividad'].tolist(), format_func=lambda x: nombres_act[x])
                
                padre = df_act[df_act['id_actividad'] == act_id].iloc[0]
                
                # Resumen de pesos usando Pandas
                df_all_subs = get_data("subactividades")
                df_sub_existentes = df_all_subs[df_all_subs['id_actividad'] == act_id]
                peso_usado = df_sub_existentes['peso'].sum()


                # --- NUEVO BLOQUE DE CONTROL ---
                valor_usado = df_sub_existentes['valor_sub'].sum()
                valor_disponible = padre['valor_total_actividad'] - valor_usado
                peso_disponible = 1.0 - peso_usado

                st.write("---")
                c1, c2, c3 = st.columns(3)
                c1.metric("Presupuesto Total", f"${padre['valor_total_actividad']:,.2f}")
                c2.metric("Peso Disponible", f"{peso_disponible:.2f}", delta=f"{peso_usado:.2f} usado", delta_color="inverse")
                c3.metric("Saldo por Distribuir", f"${valor_disponible:,.2f}")
                
                if peso_usado >= 1.0:
                    st.success("🎯 Esta actividad ya alcanzó el 100% de su peso.")
                # --- FIN DEL BLOQUE ---






                st.write(f"Peso actual: **{peso_usado:.2f} / 1.0**")

                with st.form("form_sub"):
                    sc1, sc2, sc3, sc4 = st.columns([2, 1, 1, 1])
                    sub_nombre = sc1.text_input("Nombre Subactividad")
                    sub_meta = sc2.number_input("Meta", min_value=0.0)
                    sub_unidad = sc3.text_input("Unidad")
                    sub_peso = sc4.number_input("Peso", min_value=0.0, max_value=1.0, step=0.01)
                    
                    if st.form_submit_button("➕ Vincular Subactividad"):
                        if (peso_usado + sub_peso) > 1.0001:
                            st.error("Error: Peso excedido.")
                        else:
                            sub_valor = padre['valor_total_actividad'] * sub_peso
                            df_sub = get_data("subactividades")
                            nueva_sub = pd.DataFrame([{
                                "id_sub": len(df_sub) + 1,
                                "id_actividad": act_id,
                                "nombre_subactividad": sub_nombre,
                                "valor_sub": sub_valor,
                                "meta_sub": sub_meta,
                                "unidad_medida_sub": sub_unidad,
                                "peso": sub_peso
                            }])
                            df_final = pd.concat([df_sub, nueva_sub], ignore_index=True)
                            safe_update("subactividades", df_final)
                            st.success("✅ Subactividad agregada al Excel.")
                            st.rerun()

                if not df_sub_existentes.empty:
                    st.dataframe(df_sub_existentes[['nombre_subactividad', 'meta_sub', 'unidad_medida_sub', 'peso', 'valor_sub']])


# --- TABLA DE SUBACTIVIDADES CON EDICIÓN Y ELIMINACIÓN ---
                st.write("---")
                st.subheader("📋 Subactividades Vinculadas")
                
                if not df_sub_existentes.empty:
                    # Mostrar tabla para referencia de IDs
                    st.dataframe(df_sub_existentes[['id_sub', 'nombre_subactividad', 'meta_sub', 'unidad_medida_sub', 'peso', 'valor_sub']], use_container_width=True)
                    
                    col_ed, col_del = st.columns(2)
                    
                    # --- BOTÓN EDITAR ---
                    with col_ed.expander("📝 Editar Subactividad"):
                        id_edit = st.number_input("ID de la subactividad a editar:", min_value=1, step=1, key="edit_sub_id")
                        sub_to_edit = df_sub_existentes[df_sub_existentes['id_sub'] == id_edit]
                        
                        if not sub_to_edit.empty:
                            with st.form("form_edit_sub"):
                                new_nom = st.text_input("Nuevo Nombre", value=sub_to_edit.iloc[0]['nombre_subactividad'])
                                new_meta = st.number_input("Nueva Meta", value=float(sub_to_edit.iloc[0]['meta_sub']))
                                new_peso = st.number_input("Nuevo Peso", value=float(sub_to_edit.iloc[0]['peso']), min_value=0.0, max_value=1.0)
                                
                                if st.form_submit_button("Actualizar Subactividad"):
                                    # Recalcular valor basado en el nuevo peso
                                    new_val = padre['valor_total_actividad'] * new_peso
                                    # Verificar que la suma de pesos no supere 1.0 (excluyendo el peso anterior de esta subactividad)
                                    peso_otros = df_sub_existentes[df_sub_existentes['id_sub'] != id_edit]['peso'].sum()
                                    
                                    if (peso_otros + new_peso) > 1.0001:
                                        st.error("Error: El nuevo peso hace que la suma total supere 1.0")
                                    else:
                                        df_sub_upd = get_data("subactividades")
                                        df_sub_upd.loc[df_sub_upd['id_sub'] == id_edit, ['nombre_subactividad', 'meta_sub', 'peso', 'valor_sub']] = [new_nom, new_meta, new_peso, new_val]
                                        safe_update("subactividades", df_sub_upd)
                                        st.success("✅ Subactividad actualizada en Excel")
                                        st.rerun()
                        else:
                            st.info("Ingrese un ID válido de la tabla de arriba")

                    # --- BOTÓN ELIMINAR ---
                    with col_del.expander("🗑️ Eliminar Subactividad"):
                        id_del = st.number_input("ID de la subactividad a eliminar:", min_value=1, step=1, key="del_sub_id")
                        if st.button("Confirmar Eliminación", type="primary"):
                            df_sub_del = get_data("subactividades")
                            # Filtramos para mantener todo menos el ID a borrar
                            df_sub_del = df_sub_del[df_sub_del['id_sub'] != id_del]
                            safe_update("subactividades", df_sub_del)
                            st.warning(f"⚠️ Subactividad {id_del} eliminada del Excel")
                            st.rerun()
                else:
                    st.info("No hay subactividades vinculadas a esta actividad padre.")


# --- TAB 3: ASIGNACIÓN A MUNICIPIOS ---
        with tab3:
            
            st.subheader("📍 Asignación de Presupuesto por Municipio")
            
            # Consultar subactividades uniendo DataFrames de Pandas
            df_s_raw = get_data("subactividades")
            df_a_raw = get_data("actividades_maestro")
            
            if not df_s_raw.empty and not df_a_raw.empty:
                df_sub_todas = df_s_raw.merge(df_a_raw[['id_actividad', 'nombre_actividad']], on="id_actividad")
            else:
                df_sub_todas = pd.DataFrame()


            if df_sub_todas.empty:
                st.warning("Debe configurar subactividades en la pestaña 2 antes de asignar municipios.")
            else:
                # Selector de subactividad
                sub_sel_id = st.selectbox("Seleccione Subactividad:", 
                                         df_sub_todas['id_sub'].tolist(), 
                                         format_func=lambda x: f"{df_sub_todas[df_sub_todas['id_sub']==x]['nombre_actividad'].values[0]} >> {df_sub_todas[df_sub_todas['id_sub']==x]['nombre_subactividad'].values[0]}",
                                         key="asig_muni_selector")
                
                datos_sub = df_sub_todas[df_sub_todas['id_sub'] == sub_sel_id].iloc[0]
                
           
                # Cálculos usando el DataFrame ya cargado (df_s_raw)
                df_asig_all = get_data("asignacion_municipios")
                df_asig_actual = df_asig_all[df_asig_all['id_sub'] == sub_sel_id]

                valor_gastado_muni = df_asig_actual['valor_asignado'].sum() if not df_asig_actual.empty else 0
                saldo_muni = datos_sub['valor_sub'] - valor_gastado_muni

                # Cuadros de Control
                col_m1, col_m2 = st.columns(2)


                col_m1.metric("Presupuesto Subactividad", f"${datos_sub['valor_sub']:,.2f}")
                col_m2.metric("Saldo Disponible para Municipios", f"${saldo_muni:,.2f}", delta=f"-${valor_gastado_muni:,.2f} asignado", delta_color="inverse")

                # Formulario de Registro Actualizado
                with st.form("form_municipio"):
                    m1, m2 = st.columns(2)
                    muni_nombre = m1.selectbox("Municipio de Santander", municipios_santander)
                    n_contrato = m1.text_input("Número de Contrato / Convenio")
                    n_pagos = m1.number_input("Número de Pagos (Seguimientos)", min_value=1, step=1, help="Define cuántos reportes de pago tendrá esta actividad")
                    
                    muni_valor = m2.number_input("Valor a asignar ($)", min_value=0.0, max_value=float(datos_sub['valor_sub']), step=1000.0)
                    muni_meta = m2.number_input("Meta Municipal", min_value=0.0)
                    muni_unidad = m2.text_input("Unidad de Medida Municipal (Ej: Personas, Visitas)") # <-- Nuevo campo
                    
                    if st.form_submit_button("📍 Confirmar Asignación Municipal"):
                        if muni_valor > (saldo_muni + 0.01):
                            st.error(f"Error: El valor supera el saldo disponible (${saldo_muni:,.2f})")
                        else:
                            df_asig_muni = get_data("asignacion_municipios")
                            # Generamos nuevo ID y nueva fila
                            nuevo_id_asig = 1 if df_asig_muni.empty else df_asig_muni['id_asig'].max() + 1
                            nueva_asig = pd.DataFrame([{
                                "id_asig": nuevo_id_asig,
                                "id_sub": sub_sel_id, 
                                "municipio": muni_nombre, 
                                "num_contrato": n_contrato, 
                                "num_pagos": n_pagos, 
                                "valor_asignado": muni_valor, 
                                "meta_municipal": muni_meta,
                                "unidad_medida_muni": muni_unidad # <-- Nueva columna asignada
                            }])
                            df_final_asig = pd.concat([df_asig_muni, nueva_asig], ignore_index=True)
                            safe_update("asignacion_municipios", df_final_asig)
                            st.success(f"✅ Asignación exitosa para {muni_nombre} en Excel")
                            st.rerun()

                # Tabla de Visualización Actualizada
                if not df_asig_actual.empty:
                    st.write("---")
                    st.write("**Asignaciones actuales por Municipio:**")
                    st.dataframe(df_asig_actual[['id_asig', 'municipio', 'num_contrato', 'num_pagos', 'valor_asignado', 'meta_municipal', 'unidad_medida_muni']], use_container_width=True)

                    # Opción para Eliminar
                    with st.expander("🗑️ Eliminar Asignación Municipal"):
                        id_asig_del = st.number_input("ID de asignación a eliminar:", min_value=1, step=1, key="del_asig_id")
                        if st.button("Confirmar Borrado Municipal", type="primary"):
                            df_asig_del = get_data("asignacion_municipios")
                            df_asig_del = df_asig_del[df_asig_del['id_asig'] != id_asig_del]
                            safe_update("asignacion_municipios", df_asig_del)
                            st.warning("⚠️ Asignación eliminada del Excel")
                            st.rerun()


# --- MÓDULO: EJECUCIÓN (MUNICIPIO) ---
    elif menu == "📝 Ejecución":
        if rol != "MUNICIPIO_EJECUTOR":
            st.warning("Este módulo es exclusivo para el perfil MUNICIPIO_EJECUTOR.")
        else:
            st.title("📝 Reporte de Avance Municipal")
            muni_user = st.session_state.get('muni_asignado', 'N/A') 
            

            # Carga de datos uniendo hojas de Excel con Pandas [cite: 206]
            df_a_exec = get_data("asignacion_municipios")
            df_s_exec = get_data("subactividades")
            df_m_exec = get_data("actividades_maestro")

            if not df_a_exec.empty:
                # Reemplazo del JOIN SQL por MERGE de Pandas [cite: 206]
                df_merge_exec = df_a_exec.merge(df_s_exec, on="id_sub").merge(df_m_exec, on="id_actividad")
                # Filtrar por el municipio asignado al usuario [cite: 207]
                df_mis_asig = df_merge_exec[df_merge_exec['municipio'] == muni_user]
            else:
                df_mis_asig = pd.DataFrame()



            if df_mis_asig.empty:
                st.info(f"No hay asignaciones pendientes para {muni_user}.")
            else:
                st.markdown("### 🔍 Selector de Metas Contractuales")
                
                # Filtro 1: Extraer y formatear las Actividades Maestro asignadas a este municipio
                actividades_disponibles = df_mis_asig[['id_actividad', 'nombre_actividad']].drop_duplicates()
                dict_actividades = {row['id_actividad']: f"ID {row['id_actividad']} - {row['nombre_actividad']}" for _, row in actividades_disponibles.iterrows()}
                
                sel_actividad_padre = st.selectbox(
                    "1️⃣ Filtrar por Actividad General (Maestro):", 
                    options=list(dict_actividades.keys()), 
                    format_func=lambda x: dict_actividades[x],
                    key="filtro_cascada_actividad"
                )
                
                # Filtro 2: Filtrar subactividades hijas en base a la selección del padre
                df_subs_filtradas = df_mis_asig[df_mis_asig['id_actividad'] == sel_actividad_padre]
                dict_subactividades = {row['id_asig']: f"Contrato: {row['num_contrato']} >> {row['nombre_subactividad']}" for _, row in df_subs_filtradas.iterrows()}
                
                sel_asig = st.selectbox(
                    "2️⃣ Seleccione Subactividad / Contrato Específico a Reportar:", 
                    options=list(dict_subactividades.keys()), 
                    format_func=lambda x: dict_subactividades[x],
                    key="filtro_cascada_subactividad"
                )
                
                datos = df_mis_asig[df_mis_asig['id_asig'] == sel_asig].iloc[0]
                
            
                # --- NUEVO: Extraer datos para cálculo visual ---
                meta_total = datos['meta_municipal']
                valor_total_muni = datos['valor_asignado']
                unidad_muni = datos.get('unidad_medida_muni', 'unidades')




                # Consultar último pago reportado filtrando el DataFrame
                df_pagos_all = get_data("seguimiento_pagos")
                if not df_pagos_all.empty:
                    pagos_asig = df_pagos_all[df_pagos_all['id_asig'] == sel_asig]
                    ultimo_pago = pagos_asig['num_pago_actual'].max() if not pagos_asig.empty else 0
                else:
                    ultimo_pago = 0

                siguiente_pago = ultimo_pago + 1
                
                if siguiente_pago > datos['num_pagos']:
                    st.success("✅ Todas las cuotas de pago de esta actividad han sido reportadas.")
                else:
                    st.info(f"Reportando Pago N° {siguiente_pago} de {datos['num_pagos']}") 
                    st.warning(f"📋 **Meta Total Asignada:** {meta_total} {unidad_muni} | **Presupuesto Total:** ${valor_total_muni:,.2f}")

                    # --- ENTRADA FUERA DEL FORMULARIO PARA CÁLCULO DINÁMICO ---
                    # Esto permite que el valor se actualice instantáneamente al cambiar el número
                    meta_avanc = st.number_input("Avance de Meta realizado en este periodo", min_value=0.0, step=1.0)
                    
                    val_dinamico = (meta_avanc / meta_total) * valor_total_muni if meta_total > 0 else 0.0
                    porcentaje = (meta_avanc / meta_total * 100) if meta_total > 0 else 0
                    
                    st.write(f"💰 **Valor calculado para este reporte:** :green[${val_dinamico:,.2f}]")
                    st.caption(f"*(Equivale al {porcentaje:.1f}% de la meta total municipal)*")

                    with st.form("form_reporte_muni"):
                        soporte = st.text_input("Link a carpeta de soportes (Evidencias)")
                        
                        if st.form_submit_button("Enviar a Revisión del Referente"):
                            if meta_avanc <= 0:
                                st.error("⚠️ El avance debe ser mayor a 0 para generar un cobro.")
                            else:
                                # Lectura fresca de la base de datos para evitar errores de concurrencia
                                df_pagos = get_data("seguimiento_pagos")
                                nuevo_id_seg = 1 if df_pagos.empty else df_pagos['id_seguimiento'].max() + 1
                                
                                # Consolidación del cálculo final para la base de datos
                                valor_final_pago = (meta_avanc / meta_total) * valor_total_muni
                                
                                nueva_fila_pago = pd.DataFrame([{
                                    "id_seguimiento": nuevo_id_seg,
                                    "id_asig": sel_asig,
                                    "num_pago_actual": siguiente_pago,
                                    "avance_meta": meta_avanc,
                                    "valor_calculado": valor_final_pago,
                                    "soporte_municipio": soporte,
                                    "estado": 'PENDIENTE',
                                    "fecha_registro": pd.Timestamp.now().strftime("%Y-%m-%d %H:%M")
                                }])
                                
                                df_final_pagos = pd.concat([df_pagos, nueva_fila_pago], ignore_index=True)
                                
                                # Aplicación del cambio mediante la función segura de sincronización
                                if safe_update("seguimiento_pagos", df_final_pagos):
                                    st.success(f"✅ Reporte enviado exitosamente por ${valor_final_pago:,.2f}")
                                    st.rerun()




        # --- TABLA DE SEGUIMIENTO Y ALERTAS PARA MUNICIPIO (CORREGIDO V6.3) ---
            st.write("---")
            st.subheader("📋 Historial de mis Reportes y Estados")
            
            muni_actual = st.session_state.get('muni_asignado')
            
            df_p_muni = get_data("seguimiento_pagos")
            df_a_muni = get_data("asignacion_municipios")
            df_s_muni = get_data("subactividades")
            df_m_muni = get_data("actividades_maestro")
            
            if not df_p_muni.empty and not df_a_muni.empty and not df_s_muni.empty and not df_m_muni.empty:
                # Asegurar la existencia de las columnas operativas de control en memoria
                columnas_requeridas = ['id_seguimiento', 'id_asig', 'num_pago_actual', 'valor_calculado', 'avance_meta', 'estado', 'motivo_rechazo', 'soporte_municipio']
                for col in columnas_requeridas:
                    if col not in df_p_muni.columns:
                        df_p_muni[col] = ""

                # Ensamble de traza analítica completa mediante cruzamiento de llaves con Pandas
                df_merge_muni = df_p_muni.merge(df_a_muni, on="id_asig", how="left")
                df_merge_muni = df_merge_muni.merge(df_s_muni, on="id_sub", how="left")
                df_merge_muni = df_merge_muni.merge(df_m_muni, on="id_actividad", how="left")
                
                # Aislamiento exclusivo de los registros de este municipio ejecutor
                df_muni_actual = df_merge_muni[df_merge_muni['municipio'] == muni_actual].copy()
                
                if not df_muni_actual.empty:
                    # 🔴 ALERTA 1: Renderizar rechazos activos provenientes del Supervisor General
                    df_rechazos_sup = df_muni_actual[df_muni_actual['estado'] == 'RECHAZADO']
                    if not df_rechazos_sup.empty:
                        st.error("⚠️ **ATENCIÓN DE SUPERVISIÓN:** El Supervisor ha devuelto reportes de este municipio. Revise las correcciones exigidas a continuación:")
                        for _, r_rech in df_rechazos_sup.iterrows():
                            st.warning(f"🔹 **ID Registro:** {r_rech['id_seguimiento']} | **Subactividad:** {r_rech['nombre_subactividad']} | 💬 **Motivo de Rechazo:** {r_rech['motivo_rechazo']}")
                    
                    # 🔴 ALERTA 2: Renderizar rechazos activos generados por el Referente Departamental Técnico
                    df_rechazos_ref = df_muni_actual[df_muni_actual['estado'] == 'RECHAZADO_REFERENTE']
                    if not df_rechazos_ref.empty:
                        st.error("⚠️ **ATENCIÓN TÉCNICA PIC:** El Referente Departamental ha devuelto reportes en la fase de validación. Proceda a subsanar:")
                        for _, r_rech_ref in df_rechazos_ref.iterrows():
                            st.warning(f"🔸 **ID Registro:** {r_rech_ref['id_seguimiento']} | **Subactividad:** {r_rech_ref['nombre_subactividad']} | 💬 **Observación Técnica:** {r_rech_ref['motivo_rechazo']}")
                    
                    st.markdown("#### Matriz Histórica de Trazabilidad de Pagos")
                    
                    # Filtrar las columnas clave para visualización del usuario final
                    cols_finales = ['id_seguimiento', 'id_actividad', 'nombre_actividad', 'nombre_subactividad', 'num_pago_actual', 'valor_calculado', 'avance_meta', 'estado']
                    if 'acta_referente' in df_muni_actual.columns:
                        cols_finales.append('acta_referente')
                        
                    df_visual_muni = df_muni_actual[[c for c in cols_finales if c in df_muni_actual.columns]].copy()
                    df_visual_muni.columns = ['ID Seguimiento', 'ID Actividad', 'Actividad General', 'Subactividad', 'Pago N°', 'Valor Calculado', 'Avance Meta', 'Estado'] + (['Acta Soporte'] if 'acta_referente' in df_muni_actual.columns else [])
                    
                    st.dataframe(df_visual_muni, use_container_width=True, hide_index=True)
                    
                    # --- COMPONENTE INTEGRADO DE CORRECCIÓN Y REENVÍO ---
                    with st.expander("🔄 Corregir y Reenviar Reporte Rechazado (Supervisor / Referente)"):
                        df_muni_solo_rech = df_muni_actual[df_muni_actual['estado'].isin(['RECHAZADO', 'RECHAZADO_REFERENTE'])]
                        
                        if df_muni_solo_rech.empty:
                            st.info("No se registran novedades o reportes rechazados pendientes de corrección.")
                        else:
                            opciones_rechazo_dict = {f"ID {row['id_seguimiento']} - {row['nombre_subactividad']} (Pago {row['num_pago_actual']})": row['id_seguimiento'] for _, row in df_muni_solo_rech.iterrows()}
                            sel_rechazo_label = st.selectbox("Seleccione el reporte a subsanar:", list(opciones_rechazo_dict.keys()), key="sb_muni_subsanar")
                            id_subsanar = opciones_rechazo_dict[sel_rechazo_label]
                            
                            fila_subsanar = df_muni_solo_rech[df_muni_solo_rech['id_seguimiento'] == id_subsanar].iloc[0]
                            st.markdown(f"📋 **Historial de la Alerta:** {fila_subsanar['motivo_rechazo']}")
                            
                            with st.form("form_correccion_municipio"):
                                nuevo_avance = st.number_input("Corregir avance de la meta física:", min_value=0.0, value=float(fila_subsanar['avance_meta']))
                                nuevo_soporte = st.text_input("Actualizar enlace de evidencias (URL Soportes):", value=str(fila_subsanar['soporte_municipio']))
                                
                                if st.form_submit_button("🚀 Reenviar a Validación Técnica del Referente"):
                                    if nuevo_avance <= 0:
                                        st.error("El avance físico ingresado debe ser mayor a 0.")
                                    else:
                                        df_pagos_db = get_data("seguimiento_pagos")
                                        
                                        # Recalcular el valor económico proporcional según la corrección física realizada
                                        meta_total_original = float(fila_subsanar['meta_municipal'])
                                        valor_total_original = float(fila_subsanar['valor_asignado'])
                                        valor_recalculado_pago = (nuevo_avance / meta_total_original) * valor_total_original if meta_total_original > 0 else 0.0
                                        
                                        # Actualizar el registro y resetear el estado a PENDIENTE para reiniciar el ciclo de auditoría
                                        df_pagos_db.loc[df_pagos_db['id_seguimiento'] == id_subsanar, 'avance_meta'] = nuevo_avance
                                        df_pagos_db.loc[df_pagos_db['id_seguimiento'] == id_subsanar, 'valor_calculado'] = valor_recalculado_pago
                                        df_pagos_db.loc[df_pagos_db['id_seguimiento'] == id_subsanar, 'soporte_municipio'] = str(nuevo_soporte)
                                        df_pagos_db.loc[df_pagos_db['id_seguimiento'] == id_subsanar, 'estado'] = 'PENDIENTE'
                                        df_pagos_db.loc[df_pagos_db['id_seguimiento'] == id_subsanar, 'motivo_rechazo'] = 'Reporte corregido y reenviado por el municipio'
                                        
                                        if safe_update("seguimiento_pagos", df_pagos_db):
                                            st.success("✅ Registro corregido con éxito. Ha sido enviado nuevamente al Referente Técnico.")
                                            st.rerun()
                else:
                    st.info("No hay reportes realizados aún por este municipio.")
            else:
                st.info("No hay reportes realizados aún.")





    # --- MÓDULO: REVISIÓN (REFERENTE) ---
    elif menu == "⚖️ Revisión":
        if rol not in ["REFERENTE_DEPARTAMENTAL", "SUPERVISOR"]:
            st.warning("Acceso exclusivo para personal evaluador del departamento.")
        else:
            # Cargamos las tablas necesarias
            df_p = get_data("seguimiento_pagos")
            df_a = get_data("asignacion_municipios")
            df_s = get_data("subactividades")




            # --- SUB-MÓDULO EXCLUSIVO: REFERENTE DEPARTAMENTAL (OPTIMIZADO V6 ACTAS IA + CONTRATO + OBSERVACIONES) ---
            if rol == "REFERENTE_DEPARTAMENTAL":
                st.title("⚖️ Validación de Reportes y Generación de Actas de Conformidad")
                
                # Ensamble integral del esquema analítico (Finanzas + Salud Pública) [cite: 222]
                if not df_p.empty and not df_a.empty and not df_s.empty:
                    df_merge_ref = df_p.merge(df_a, on="id_asig").merge(df_s, on="id_sub")
                    if not df_act_raw.empty:
                        df_merge_ref = df_merge_ref.merge(df_act_raw[['id_actividad', 'nombre_actividad', 'programa_responsable']], on="id_actividad", how="left")
                else:
                    df_merge_ref = pd.DataFrame()

                # Segmentación de UX: Gestión operativa vs Auditoría de Actas por Periodo [cite: 223]
                tab_gestion, tab_acta_ia = st.tabs(["📥 Validar Reportes Entrantes", "📄 Generador de Actas Certificadas (IA)"])

                with tab_gestion:
                    st.subheader("Reportes Municipales Pendientes de Revisión")
                    if not df_merge_ref.empty:
                        df_pendientes = df_merge_ref[df_merge_ref['estado'] == 'PENDIENTE'].copy()
                    else:
                        df_pendientes = pd.DataFrame()

                    if df_pendientes.empty:
                        st.info("No se registran reportes municipales en estado PENDIENTE.")
                    else:
                        df_visual_pend = df_pendientes[['id_seguimiento', 'municipio', 'num_contrato', 'nombre_subactividad', 'num_pago_actual', 'valor_calculado', 'soporte_municipio']].copy()
                        df_visual_pend.columns = ['ID Seguimiento', 'Municipio', 'Contrato', 'Subactividad', 'Pago N°', 'Valor a Validar', 'URL Soporte']
                        st.dataframe(df_visual_pend, use_container_width=True, hide_index=True)
                        
                        with st.form("form_revision_referente"):
                            id_rev = st.number_input("ID de Seguimiento a Validar / Rechazar:", min_value=1, step=1)
                            decision_ref = st.radio("Dictamen de la Revisión:", ["Aprobar Validación Técnica", "Rechazar y Devolver al Municipio"])
                            acta_link = st.text_input("Enlace General al Repositorio de Evidencias / Acta (Opcional)")
                            obs_tecnicas = st.text_area("✍️ Observaciones de la Revisión / Motivo del Rechazo (Obligatorio si rechaza):")
                            
                            if st.form_submit_button("Confirmar Dictamen Técnico 🚀"):
                                if decision_ref == "Rechazar y Devolver al Municipio" and not obs_tecnicas.strip():
                                    st.error("❌ Debe ingresar el motivo del rechazo en el campo de observaciones.")
                                else:
                                    df_rev = get_data("seguimiento_pagos")
                                    columnas_texto = ['estado', 'acta_referente', 'referente_aprobador', 'observaciones_referente', 'motivo_rechazo']
                                    for col in columnas_texto:
                                        if col in df_rev.columns:
                                            df_rev[col] = df_rev[col].astype(str).replace(['nan', 'None', '<NA>'], '')
                                        else:
                                            df_rev[col] = ''
                                        
                                    if decision_ref == "Aprobar Validación Técnica":
                                        df_rev.loc[df_rev['id_seguimiento'] == id_rev, 'estado'] = 'REVISADO_REFERENTE'
                                        df_rev.loc[df_rev['id_seguimiento'] == id_rev, 'motivo_rechazo'] = 'Avalado por el referente'
                                    else:
                                        df_rev.loc[df_rev['id_seguimiento'] == id_rev, 'estado'] = 'RECHAZADO_REFERENTE'
                                        df_rev.loc[df_rev['id_seguimiento'] == id_rev, 'motivo_rechazo'] = f"RECHAZO REFERENTE: {str(obs_tecnicas)}"
                                        
                                    df_rev.loc[df_rev['id_seguimiento'] == id_rev, 'acta_referente'] = str(acta_link)
                                    df_rev.loc[df_rev['id_seguimiento'] == id_rev, 'observaciones_referente'] = str(obs_tecnicas)
                                    df_rev.loc[df_rev['id_seguimiento'] == id_rev, 'referente_aprobador'] = str(st.session_state['user'])

                                    if safe_update("seguimiento_pagos", df_rev):
                                        st.success(f"✅ Dictamen del reporte ID {id_rev} sincronizado exitosamente.")
                                        st.rerun()


                with tab_acta_ia:
                    st.subheader("Filtros de Consolidación para el Acta Formal")
                    st.caption("Seleccione el territorio y periodo para compilar las actividades revisadas y generar el análisis descriptivo.")

                    if df_merge_ref.empty:
                        st.warning("No existen registros consolidados en la base de datos para estructurar actas.")
                    else:
                        c_f1, c_f2 = st.columns(2)
                        municipios_activos = sorted(df_merge_ref['municipio'].dropna().unique().tolist())
                        muni_seleccionado = c_f1.selectbox("📍 Seleccione el Municipio:", municipios_activos, key="ref_muni_acta")
                        
                        periodos_disponibles = sorted(df_merge_ref['num_pago_actual'].dropna().unique().astype(int).tolist())
                        periodo_seleccionado = c_f2.selectbox("📆 Seleccione el Periodo de Pago / Cuota:", periodos_disponibles, key="ref_pago_acta")

                        # Aislamiento matricial de las actividades revisadas por el referente (estado REVISADO_REFERENTE o ACEPTADA)
                        df_actas_filtrado = df_merge_ref[
                            (df_merge_ref['municipio'] == muni_seleccionado) & 
                            (df_merge_ref['num_pago_actual'].astype(int) == int(periodo_seleccionado)) &
                            (df_merge_ref['estado'].isin(['REVISADO_REFERENTE', 'ACEPTADA']))
                        ].copy()

                        if df_actas_filtrado.empty:
                            st.info(f"ℹ️ No se han encontrado actividades con validación técnica para **{muni_seleccionado}** en el **Periodo {periodo_seleccionado}**.")
                        else:
                            st.success(f"📋 Se detectaron **{len(df_actas_filtrado)}** subactividades procesadas para el acta.")
                            
                            # --- EXTRACCIÓN DINÁMICA DEL NÚMERO DE CONTRATO ---
                            contrato_municipio = str(df_actas_filtrado['num_contrato'].iloc[0]) if 'num_contrato' in df_actas_filtrado.columns and not df_actas_filtrado['num_contrato'].empty else "POR DEFINIR"
                            
                            st.info(f"📄 **Información del Contrato Localizado:** {contrato_municipio}")
                            
                            # Normalizar la columna de observaciones del referente para prevenir fallas visuales
                            if 'observaciones_referente' not in df_actas_filtrado.columns:
                                df_actas_filtrado['observaciones_referente'] = "Sin observaciones registradas"
                            else:
                                df_actas_filtrado['observaciones_referente'] = df_actas_filtrado['observaciones_referente'].fillna("Sin observaciones")

                            df_resumen_tabla = df_actas_filtrado[['id_seguimiento', 'nombre_subactividad', 'meta_municipal', 'avance_meta', 'valor_calculado', 'observaciones_referente']].copy()
                            df_resumen_tabla.columns = ['ID Seguimiento', 'Subactividad', 'Meta Municipal', 'Avance Físico', 'Valor Devengado ($)', 'Observaciones del Referente']
                            st.dataframe(df_resumen_tabla, use_container_width=True, hide_index=True)

                            # Cálculos agregados científicos de datos y financieros
                            total_financiero_periodo = df_actas_filtrado['valor_calculado'].sum()
                            total_metas_programadas = df_actas_filtrado['meta_municipal'].sum()
                            total_metas_ejecutadas = df_actas_filtrado['avance_meta'].sum()
                            eficiencia_operativa_acta = (total_metas_ejecutadas / total_metas_programadas * 100) if total_metas_programadas > 0 else 0
                            
                            # Estructuración de cadena sintética contextual para el prompt de la IA incluyendo observaciones del referente
                            lineas_actividades = []
                            for idx, row in df_actas_filtrado.iterrows():
                                programa_txt = row.get('programa_responsable', 'Salud Pública General')
                                obs_fila = row.get('observaciones_referente', 'Sin observaciones')
                                lineas_actividades.append(
                                    f"- Actividad: {row['nombre_subactividad']} | Componente: {programa_txt} | "
                                    f"Meta Programada: {row['meta_municipal']} | Avance Realizado: {row['avance_meta']} | "
                                    f"Valor Proporcional Líquido: ${row['valor_calculado']:,.2f} | "
                                    f"Dictamen Técnico del Referente: {obs_fila}"
                                )
                            bloque_actividades_contexto = "\n".join(lineas_actividades)

                            if st.button("🤖 Compilar y Redactar Acta con IA Experta", type="primary", key="btn_generar_acta_ia"):
                                with st.spinner("Conectando con el motor analítico de salud pública..."):
                                    
                                    prompt_acta = f"""
                                    Actúa como un Referente Departamental de Salud Pública, Auditor Técnico de Cuentas y Científico de Datos.
                                    Escribe el cuerpo analítico de una CERTIFICACIÓN / ACTA DE CONFORMIDAD TÉCNICA Y FINANCIERA para el Plan de Intervenciones Colectivas (PIC).
                                    
                                    DATOS DE CONTROL GEOGRÁFICO, TEMPORAL Y CONTRACTUAL:
                                    - Municipio Auditado: {muni_seleccionado}
                                    - Número de Contrato Estatal: {contrato_municipio}
                                    - Cuota de Pago / Periodo: Periodo N° {periodo_seleccionado}
                                    - Profesional Evaluador Emisor: {st.session_state['user']}
                                    
                                    MÉTRICAS AGREGADAS DEL TERRITORIO:
                                    - Presupuesto Total Devengado en este Periodo: ${total_financiero_periodo:,.2f}
                                    - Sumatoria de Metas Físicas Programadas: {total_metas_programadas}
                                    - Sumatoria de Avances Físicos Reportados: {total_metas_ejecutadas}
                                    - Índice de Eficiencia Física Global del Periodo: {eficiencia_operativa_acta:.2f}%
                                    
                                    DESGLOSE DE ACTIVIDADES EVALUADAS TÉCNICAMENTE Y OBSERVACIONES DEL REFERENTE:
                                    {bloque_actividades_contexto}
                                    
                                    REQUISITOS FORMALES DE REDACCIÓN DEL DICTAMEN:
                                    Escribe el documento con tono corporativo, pericial, riguroso y formal. Debe hacer mención expresa al Contrato N° {contrato_municipio}, sintetizar e integrar obligatoriamente el hallazgo plasmado en el 'Dictamen Técnico del Referente' de cada subactividad y estructurar:
                                    1. ANÁLISIS FINANCIERO DE LA VIGENCIA (Concepto técnico sobre el valor devengado, concordancia con el presupuesto asignado y proporcionalidad del cobro bajo el marco legal del contrato).
                                    2. ANÁLISIS OPERATIVO Y SANITARIO (Evaluación epidemiológica del cumplimiento de metas físicas, justificación del avance en campo y suficiencia de las evidencias e incorporando las observaciones técnicas recopiladas).
                                    3. OBSERVACIONES O RECOMENDACIONES TÉCNICAS (Mínimo 2 sugerencias viables de optimización de procesos basadas en los datos expuestos para el siguiente periodo).
                                    """

                                    try:
                                        import requests
                                        import json

                                        api_key_ref = st.secrets["GEMINI_API_KEY"]
                                        url_api_ref = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent?key={api_key_ref}"
                                        
                                        headers_ref = {"Content-Type": "application/json"}
                                        payload_ref = {
                                            "contents": [{
                                                "parts": [{"text": prompt_acta}]
                                            }]
                                        }

                                        res_ref = requests.post(url_api_ref, headers=headers_ref, data=json.dumps(payload_ref))
                                        
                                        if res_ref.status_code == 200:
                                            datos_json_ref = res_ref.json()
                                            acta_ia_texto = datos_json_ref['candidates'][0]['content']['parts'][0]['text']
                                        else:
                                            raise Exception(f"Fallo en comunicación externa (Código: {res_ref.status_code})")
                                            
                                    except Exception as e:
                                        acta_ia_texto = f"ANÁLISIS TÉCNICO-FINANCIERO CONTRACTUAL:\n\nEl municipio de {muni_seleccionado}, bajo el amparo del Contrato N° {contrato_municipio}, presenta una ejecución financiera certificada de ${total_financiero_periodo:,.2f} durante el Periodo N° {periodo_seleccionado}.\n\nSe evidencia un índice de eficiencia operativa promedio del {eficiencia_operativa_acta:.2f}% en el cumplimiento de las metas físicas estipuladas. Las observaciones del área técnica reafirman que las evidencias técnicas anexadas en el repositorio digital institucional cumplen formalmente con los lineamientos del manual de supervisión departamental del PIC."

                                    st.markdown("### 📄 Cuerpo de Acta Generado por IA")
                                    st.markdown(acta_ia_texto)

                                    # --- CONSTRUCCIÓN DINÁMICA DEL DOCUMENTO WORD (.DOCX) ---
                                    from docx import Document
                                    from docx.shared import Inches, Pt
                                    import io

                                    doc_acta = Document()
                                    
                                    # Márgenes de auditoría estándar de 1 pulgada
                                    for section in doc_acta.sections:
                                        section.top_margin = Inches(1)
                                        section.bottom_margin = Inches(1)
                                        section.left_margin = Inches(1)
                                        section.right_margin = Inches(1)

                                    # Encabezado Formal del Acta
                                    p_title = doc_acta.add_paragraph()
                                    run_title = p_title.add_run("ACTA DE CONFORMIDAD Y CERTIFICACIÓN TÉCNICA - PIC")
                                    run_title.bold = True
                                    run_title.font.size = Pt(14)
                                    run_title.font.name = 'Arial'
                                    p_title.alignment = 1

                                    # Bloque de Control de Metadatos Administrativos
                                    doc_acta.add_paragraph(f"Municipio Beneficiario: {muni_seleccionado}")
                                    doc_acta.add_paragraph(f"Número de Contrato Estatal: {contrato_municipio}")
                                    doc_acta.add_paragraph(f"Periodo de Pago / Cuota Evaluada: Periodo N° {periodo_seleccionado}")
                                    doc_acta.add_paragraph(f"Fecha de Certificación: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')}")
                                    doc_acta.add_paragraph(f"Profesional Referente Evaluador: {st.session_state['user']}")
                                    doc_acta.add_paragraph("------------------------------------------------------------------------------------------------------------------------")

                                    # Tabla Estructural de Datos Cuantitativos y Cualitativos
                                    doc_acta.add_heading('1. Balance Cuantitativo y Observaciones de Actividades Revisadas', level=2)
                                    tabla_ref = doc_acta.add_table(rows=1, cols=5)
                                    tabla_ref.style = 'Light Shading Accent 1'
                                    hdr_cells_ref = tabla_ref.rows[0].cells
                                    hdr_cells_ref[0].text = 'Subactividad Evaluada'
                                    hdr_cells_ref[1].text = 'Meta Prog.'
                                    hdr_cells_ref[2].text = 'Avance Físico'
                                    hdr_cells_ref[3].text = 'Monto Proporcional'
                                    hdr_cells_ref[4].text = 'Observación Técnica del Referente'

                                    for _, fila_acta in df_actas_filtrado.iterrows():
                                        row_c = tabla_ref.add_row().cells
                                        row_c[0].text = str(fila_acta['nombre_subactividad'])
                                        row_c[1].text = str(fila_acta['meta_municipal'])
                                        row_c[2].text = str(fila_acta['avance_meta'])
                                        row_c[3].text = f"${fila_acta['valor_calculado']:,.2f}"
                                        row_c[4].text = str(fila_acta['observaciones_referente'])

                                    p_totales = doc_acta.add_paragraph()
                                    p_totales.add_run(f"\nTOTAL FINANCIERO CERTIFICADO EN EL PERIODO: ${total_financiero_periodo:,.2f}\n").bold = True
                                    p_totales.add_run(f"EFICIENCIA OPERATIVA TERRITORIAL PROMEDIO: {eficiencia_operativa_acta:.2f}%").bold = True

                                    # Inyección del dictamen analítico emitido por la IA
                                    doc_acta.add_heading('2. Dictamen de Validación Técnica y Análisis de Impacto', level=2)
                                    doc_acta.add_paragraph(acta_ia_texto)

                                    # Bloque Formal de Firmas de Cierre
                                    doc_acta.add_paragraph("\n\n\n_________________________________________\n"
                                                           "Firma y Aval de Conformidad\n"
                                                           f"Referente Departamental: {st.session_state['user']}")

                                    # Conversión del buffer binario para descarga asíncrona
                                    bio_acta = io.BytesIO()
                                    doc_acta.save(bio_acta)
                                    bio_acta.seek(0)

                                    st.download_button(
                                        label="📥 Descargar Acta de Conformidad Oficial (.docx)",
                                        data=bio_acta,
                                        file_name=f"Acta_Conformidad_PIC_{muni_seleccionado}_Contrato_{contrato_municipio}_Pago_{periodo_seleccionado}.docx",
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                        key="btn_download_acta_word"
                                    )





            # --- SUB-MÓDULO EXCLUSIVO: SUPERVISOR DEPARTAMENTAL ---
            elif rol == "SUPERVISOR":
                st.title("👁️ Panel de Supervisión y Dictamen Final")
                st.write("Audite los reportes que ya cuentan con el aval técnico del Referente asignado.")

                if not df_p.empty:
                    df_merge = df_p.merge(df_a, on="id_asig").merge(df_s, on="id_sub")
                    
                    # Normalización defensiva de columnas opcionales
                    for col in ['referente_aprobador', 'acta_referente', 'motivo_rechazo', 'soporte_municipio']:
                        if col not in df_merge.columns:
                            df_merge[col] = ""

                    # El supervisor evalúa lo que el referente ya revisó
                    df_pendientes_super = df_merge[df_merge['estado'] == 'REVISADO_REFERENTE']
                else:
                    df_pendientes_super = pd.DataFrame()

                if df_pendientes_super.empty:
                    st.success("✅ No hay actividades pendientes por evaluar de su parte.")
                else:
                    st.subheader(f"📋 Reportes listos para Dictamen Final ({len(df_pendientes_super)})")
                    
                    # Diccionario indexado para el selector visual
                    opciones_super = {
                        f"ID {row['id_seguimiento']} - {row['municipio']} (Pago {row['num_pago_actual']})": row['id_seguimiento']
                        for _, row in df_pendientes_super.iterrows()
                    }
                    
                    sel_texto = st.selectbox("Seleccione el registro a auditar:", list(opciones_super.keys()))
                    id_evaluar = opciones_super[sel_texto]
                    
                    fila_sel = df_pendientes_super[df_pendientes_super['id_seguimiento'] == id_evaluar].iloc[0]
                    
                    # Ficha de Auditoría Visual
                    with st.container(border=True):
                        c_aud1, c_aud2 = st.columns(2)
                        with c_aud1:
                            st.write(f"**Municipio:** {fila_sel['municipio']}")
                            st.write(f"**Actividad:** {fila_sel['nombre_subactividad']}")
                            st.write(f"**Monto Proporcional:** ${fila_sel['valor_calculado']:,.2f}")
                            st.write(f"**💬 Observaciones del Referente:** {fila_sel.get('observaciones_referente', 'Sin observaciones')}")
                        with c_aud2:
                            st.write(f"**🧑‍💼 Avalado por Referente:** :blue[{fila_sel['referente_aprobador']}]")
                            st.write(f"🔗 [Ver Evidencias del Municipio]({fila_sel['soporte_municipio']})")
                            st.write(f"📄 [Ver Acta Adjunta del Referente]({fila_sel['acta_referente']})")



                    # Formulario de Dictamen Final con Auditoría y Lista de Chequeo Administrativa
                    with st.form("form_dictamen_supervisor"):
                        st.markdown("### 📋 Lista de Chequeo de Seguimiento Administrativo")
                        st.caption("Diligencie el cumplimiento normativo aplicable al periodo bajo examen.")
                        
                        opciones_chk = ["Si", "No", "No Aplica"]
                        
                        c_chk1, c_chk2 = st.columns(2)
                        val_plan = c_chk1.selectbox("1. ¿El contratante entrega Plan de Trabajo concertado?", opciones_chk)
                        val_crono = c_chk1.selectbox("2. ¿Entrega Cronograma mensualizado de ejecución?", opciones_chk)
                        val_pers = c_chk1.selectbox("3. ¿Cuenta con la totalidad del personal idóneo solicitado?", opciones_chk)
                        val_segsoc = c_chk1.selectbox("4. ¿Realiza y anexa el pago de Seguridad Social de las planillas?", opciones_chk)
                        
                        val_iparcial = c_chk2.selectbox("5. ¿Se entrega el Informe Parcial técnico correspondiente?", opciones_chk)
                        val_ifinal = c_chk2.selectbox("6. ¿Se entrega el Informe Final consolidado del proceso?", opciones_chk)
                        val_polizas = c_chk2.selectbox("7. ¿Las pólizas de cumplimiento contractual se encuentran vigentes?", opciones_chk)
                        
                        st.markdown("---")
                        dictamen = st.radio("Resolución de la Supervisión:", ["Aprobar Actividad (ACEPTADA)", "Rechazar y Devolver al Municipio"])
                        motivo = st.text_area("Motivo del rechazo / Observaciones (Obligatorio en caso de rechazo):")
                        
                        if st.form_submit_button("Confirmar y Sincronizar Dictamen"):
                            df_pagos_master = get_data("seguimiento_pagos")
                            
                            # Asegurar columnas en la matriz estructural de persistencia
                            columnas_auditoria = [
                                'estado', 'motivo_rechazo', 'supervisor_aprobador',
                                'chk_plan_trabajo', 'chk_cronograma', 'chk_personal', 
                                'chk_seg_social', 'chk_inf_parcial', 'chk_inf_final', 'chk_polizas'
                            ]
                            for c in columnas_auditoria:
                                if c not in df_pagos_master.columns:
                                    df_pagos_master[c] = ""

                            if dictamen == "Aprobar Actividad (ACEPTADA)":
                                # --- SOLUCIÓN ANTICASH: Conversión explícita y preventiva de columnas de texto de auditoría ---
                                columnas_auditoria_texto = [
                                    'estado', 'motivo_rechazo', 'supervisor_aprobador',
                                    'chk_plan_trabajo', 'chk_cronograma', 'chk_personal', 
                                    'chk_seg_social', 'chk_inf_parcial', 'chk_inf_final', 'chk_polizas'
                                ]
                                for col in columnas_auditoria_texto:
                                    if col in df_pagos_master.columns:
                                        df_pagos_master[col] = df_pagos_master[col].astype(str).replace(['nan', 'None', '<NA>'], '')
                                    else:
                                        df_pagos_master[col] = ''
                                
                                # Asignación de variables operativas, financieras y la nueva lista de chequeo administrativa (Alineadas de forma segura)
                                df_pagos_master.loc[df_pagos_master['id_seguimiento'] == id_evaluar, 'estado'] = 'ACEPTADA'
                                df_pagos_master.loc[df_pagos_master['id_seguimiento'] == id_evaluar, 'motivo_rechazo'] = 'Aprobado sin novedades'
                                df_pagos_master.loc[df_pagos_master['id_seguimiento'] == id_evaluar, 'supervisor_aprobador'] = str(st.session_state['user'])
                                
                                df_pagos_master.loc[df_pagos_master['id_seguimiento'] == id_evaluar, 'chk_plan_trabajo'] = str(val_plan)
                                df_pagos_master.loc[df_pagos_master['id_seguimiento'] == id_evaluar, 'chk_cronograma'] = str(val_crono)
                                df_pagos_master.loc[df_pagos_master['id_seguimiento'] == id_evaluar, 'chk_personal'] = str(val_pers)
                                df_pagos_master.loc[df_pagos_master['id_seguimiento'] == id_evaluar, 'chk_seg_social'] = str(val_segsoc)
                                df_pagos_master.loc[df_pagos_master['id_seguimiento'] == id_evaluar, 'chk_inf_parcial'] = str(val_iparcial)
                                df_pagos_master.loc[df_pagos_master['id_seguimiento'] == id_evaluar, 'chk_inf_final'] = str(val_ifinal)
                                df_pagos_master.loc[df_pagos_master['id_seguimiento'] == id_evaluar, 'chk_polizas'] = str(val_polizas)

                                
                                if safe_update("seguimiento_pagos", df_pagos_master):
                                    st.success(f"✅ El pago ID {id_evaluar} fue marcado como ACEPTADA con su trazabilidad administrativa.")
                                    st.rerun()



                            else:
                                if not motivo.strip():
                                    st.error("❌ Error estructural: Debe ingresar el motivo del rechazo para poder devolver el trámite.")
                                else:
                                    # Forzamos la conversión a tipo string para evitar conflictos de dtypes de Pandas
                                    df_pagos_master['estado'] = df_pagos_master['estado'].astype(str)
                                    df_pagos_master['motivo_rechazo'] = df_pagos_master['motivo_rechazo'].astype(str)
                                    
                                    # El estado regresa a PENDIENTE, reiniciando el flujo para el Municipio de forma segura
                                    df_pagos_master.loc[df_pagos_master['id_seguimiento'] == id_evaluar, 'estado'] = 'PENDIENTE'
                                    df_pagos_master.loc[df_pagos_master['id_seguimiento'] == id_evaluar, 'motivo_rechazo'] = motivo
                                    
                                    if safe_update("seguimiento_pagos", df_pagos_master):
                                        st.warning(f"⚠️ Reporte devuelto al municipio. Estado restablecido a PENDIENTE.")
                                        st.rerun()
# ==============================================================================
# GENERADOR DE INFORMES DE EJECUCIÓN: INFORME TECNICO ADMINISTRATIVO Y FINANCIERO
# ==============================================================================
            st.write("---")
            st.header("📄 Generador de Informes Técnicos, Administrativos y Financieros")
            st.caption("Ecosistema analítico avanzado para la supervisión integrada, auditoría presupuestal y trazabilidad de giros PIC.")

            # 1. Filtros de generación con extracción dinámica de municipios para evitar KeyError
            df_a_inf_init = get_data("asignacion_municipios")
            if not df_a_inf_init.empty and 'municipio' in df_a_inf_init.columns:
                lista_municipios_dinamica = sorted(df_a_inf_init['municipio'].dropna().unique().tolist())
            else:
                lista_municipios_dinamica = ["SANTANDER"]

            c_inf1, c_inf2 = st.columns(2)
            tipo_informe = c_inf1.selectbox("Ámbito del Informe Cobertura", ["DEPARTAMENTAL", "MUNICIPIO ESPECÍFICO"])
            
            muni_filtro = None
            if tipo_informe == "MUNICIPIO ESPECÍFICO":
                muni_filtro = c_inf2.selectbox("Seleccione el Municipio a Evaluar", lista_municipios_dinamica)

            if 'ultimo_informe_tiempo' not in st.session_state:
                st.session_state['ultimo_informe_tiempo'] = None

            if st.button("📊 Construir e Invocar Reporte Experto PIC"):
                import time
                ahora = time.time()
                tiempo_limite = 300  # 5 minutos de ventana de seguridad
                
                if st.session_state['ultimo_informe_tiempo'] is not None:
                    tiempo_transcurrido = ahora - st.session_state['ultimo_informe_tiempo']
                    if tiempo_transcurrido < tiempo_limite:
                        minutos_restantes = int((tiempo_limite - tiempo_transcurrido) // 60)
                        segundos_restantes = int((tiempo_limite - tiempo_transcurrido) % 60)
                        st.error(f"🛑 Control de Seguridad: Ha generado un reporte recientemente. Por favor espere {minutos_restantes} min y {segundos_restantes} seg antes de solicitar otro análisis a la IA.")
                        st.stop()

                with st.spinner("Procesando balances financieros, analizando flujo de pagos y estructurando informe técnico..."):
                    
                    df_p_inf = get_data("seguimiento_pagos")
                    df_a_inf = get_data("asignacion_municipios")
                    df_s_inf = get_data("subactividades")
                    df_m_inf = get_data("actividades_maestro")
                    df_sec_inf = get_data("secuencia")
                    
                    if df_p_inf.empty or df_a_inf.empty or df_s_inf.empty or df_m_inf.empty:
                        st.error("Datos insuficientes en el sistema para consolidar un informe.")
                    else:
                        # Ensamble matricial relacional de Pandas
                        df_completo = df_p_inf.merge(df_a_inf, on="id_asig").merge(df_s_inf, on="id_sub").merge(df_m_inf, on="id_actividad", suffixes=('', '_maestro'))
                        
                        # Limpieza defensiva de tipos numéricos en secuencia de pagos
                        if df_sec_inf is not None and not df_sec_inf.empty:
                            df_sec_clean = df_sec_inf.copy()
                            df_sec_clean['valor_cp'] = pd.to_numeric(df_sec_clean['valor_cp'], errors='coerce').fillna(0.0)
                            df_sec_clean['total_pagado_oc'] = pd.to_numeric(df_sec_clean['total_pagado_oc'], errors='coerce').fillna(0.0)
                            df_sec_clean['saldo_cp'] = pd.to_numeric(df_sec_clean['saldo_cp'], errors='coerce').fillna(0.0)
                        else:
                            df_sec_clean = pd.DataFrame(columns=['municipio', 'valor_cp', 'total_pagado_oc', 'saldo_cp'])

                        # 2. BIFURCACIÓN DE INGENIERÍA DE DATOS Y MÉTRICAS SEGÚN FILTRO SELECCIONADO
                        # 2. BIFURCACIÓN DE INGENIERÍA DE DATOS Y MÉTRICAS SEGÚN FILTRO SELECCIONADO (V6.7 OPERATIVO)
                        if tipo_informe == "DEPARTAMENTAL":
                            df_filtrado = df_completo.copy() [cite: 2067]
                            nombre_informe_titulo = "INFORME TECNICO ADMINISTRATIVO Y FINANCIERO - CONSOLIDADO DEPARTAMENTAL" [cite: 2067]
                            
                            # Estadísticas Financieras Globales
                            macro_total_pic = df_m_inf['valor_total_actividad'].sum() [cite: 2068]
                            macro_total_asig = df_a_inf['valor_asignado'].astype(float).sum() [cite: 2068]
                            macro_ejecutado_pago = df_filtrado[df_filtrado['estado'].isin(['ACEPTADA', 'REVISADO_REFERENTE'])]['valor_calculado'].sum() [cite: 2069]
                            macro_pagado_efectivo = df_sec_clean['total_pagado_oc'].sum() [cite: 2069]
                            macro_reserva_disponible = macro_total_pic - macro_total_asig [cite: 2069]
                            
                            # --- INGENIERÍA OPERATIVA DEPARTAMENTAL ADVANCED ---
                            # Conteo exacto de municipios con ejecuciones activas frente al total territorial
                            municipios_con_ejecucion = df_filtrado[df_filtrado['estado'].isin(['ACEPTADA', 'REVISADO_REFERENTE'])]['municipio'].nunique()
                            tasa_cobertura_municipios = (municipios_con_ejecucion / 87 * 100)
                            
                            # Análisis operativo de la columna meta_global (Actividades Maestro)
                            meta_global_pic_nominal = df_m_inf['meta_global'].astype(float).sum()
                            avance_global_pic_real = df_filtrado[df_filtrado['estado'].isin(['ACEPTADA', 'REVISADO_REFERENTE'])]['avance_meta'].sum()
                            eficiencia_fisica_global = (avance_global_pic_real / meta_global_pic_nominal * 100) if meta_global_pic_nominal > 0 else 0.0
                            
                            total_actividades_evaluadas = df_filtrado['id_seguimiento'].nunique() [cite: 2070]
                            actividades_aceptadas = df_filtrado[df_filtrado['estado'] == 'ACEPTADA']['id_seguimiento'].nunique() [cite: 2070]
                            porcentaje_operativo = eficiencia_fisica_global
                            porcentaje_ejecucion = (macro_ejecutado_pago / macro_total_asig * 100) if macro_total_asig > 0 else 0.0 [cite: 2071]
                            promedio_pago = df_filtrado['valor_calculado'].mean() if not df_filtrado.empty else 0.0 [cite: 2072]
                            
                            bloque_financiero_contexto = f"""
                            - Ámbito de Análisis: Nivel Macroeconómico Departamental (Santander) [cite: 2072, 2073]
                            - Total del Valor del Techo Global PIC: ${macro_total_pic:,.2f} [cite: 2073]
                            - Total Asignado a los Municipios: ${macro_total_asig:,.2f} [cite: 2073]
                            - Total Ejecutado Aceptado para Pago: ${macro_ejecutado_pago:,.2f} [cite: 2074]
                            - Total Pagado Efectivo (Desembolso de Caja OC): ${macro_pagado_efectivo:,.2f} [cite: 2074]
                            - Saldo Libre de Bolsa Central: ${macro_reserva_disponible:,.2f} [cite: 2074]
                            """
                            
                            bloque_operativo_contexto = f"""
                            - Análisis Técnico-Operativo y de Cobertura Sanitaria: Departamental
                            - Total de Municipios de Santander con Ejecución Activa de Actividades: {municipios_con_ejecucion} de 87 entidades ({tasa_cobertura_municipios:.1f}%).
                            - Sumatoria de la Métrica 'Meta Global' PIC Parametrizada (Techo de Metas): {meta_global_pic_nominal:,.1f} unidades.
                            - Sumatoria de Avances Reales Consolidados en el Departamento: {avance_global_pic_real:,.1f} unidades.
                            - Índice General de Cumplimiento Técnico de la Meta Global PIC: {eficiencia_fisica_global:.2f}%
                            - Total de Líneas de Operación Reportadas: {total_actividades_evaluadas} registros analizados. [cite: 2075]
                            """
                            
                        else:
                            df_filtrado = df_completo[df_completo['municipio'] == muni_filtro].copy() [cite: 2078]
                            nombre_informe_titulo = f"INFORME TECNICO ADMINISTRATIVO Y FINANCIERO - MUNICIPIO: {muni_filtro.upper()}" [cite: 2078]
                            
                            if df_filtrado.empty: 
                                st.warning(f"⚠️ No se encontraron registros de ejecución para el municipio {muni_filtro}.")
                                st.stop() 
                                
                            # Estadísticas Financieras Locales de Precisión
                            muni_total_asig = df_a_inf[df_a_inf['municipio'] == muni_filtro]['valor_asignado'].astype(float).sum() [cite: 2080]
                            muni_ejecutado_pago = df_filtrado[df_filtrado['estado'].isin(['ACEPTADA', 'REVISADO_REFERENTE'])]['valor_calculado'].sum() [cite: 2081]
                            
                            df_sec_muni = df_sec_clean[df_sec_clean['municipio'] == muni_filtro] [cite: 2081]
                            muni_pagado_efectivo = df_sec_muni['total_pagado_oc'].sum() [cite: 2082]
                            muni_saldo_reserva = df_sec_muni['saldo_cp'].sum() [cite: 2082]
                            
                            # --- INGENIERÍA OPERATIVA MUNICIPAL ADVANCED ---
                            # Análisis cruzado de metas de subactividades y ejecución real local
                            muni_meta_sub_nominal = df_filtrado['meta_sub'].astype(float).sum()
                            meta_programada = df_filtrado['meta_municipal'].sum() [cite: 2083]
                            meta_avanzada = df_filtrado['avance_meta'].sum() [cite: 2083]
                            
                            porcentaje_operativo = (meta_avanzada / meta_programada * 100) if meta_programada > 0 else 0.0 [cite: 2083, 2084]
                            porcentaje_ejecucion = (muni_ejecutado_pago / muni_total_asig * 100) if muni_total_asig > 0 else 0.0 [cite: 2084]
                            promedio_pago = df_filtrado['valor_calculado'].mean() if not df_filtrado.empty else 0.0 [cite: 2084]
                            
                            # Estructuración analítica de la matriz interna de metas locales cruzada

                            # Estructuración de la matriz interna de metas locales cruzada
                            # ✅ CÓVIGO CORREGIDO CON SANGRÍA DE PRODUCCIÓN EXACTA (Usa barra espaciadora):
                            # Estructuración de la matriz interna de metas locales cruzada
                            lineas_desglose = []
                            for _, row in df_filtrado.iterrows():
                                lineas_desglose.append(
                                    f"  * Actividad General Maestro ID {row['id_actividad']} ({row['nombre_actividad'][:25]}) -> Subactividad: {row['nombre_subactividad']} | Meta Teórica Subactividad: {row['meta_sub']} | Meta Local Asignada: {row['meta_municipal']} | Avance Real de Meta en Campo: {row['avance_meta']} | Porcentaje de Cumplimiento Local: {(float(row['avance_meta'])/float(row['meta_municipal'])*100 if float(row['meta_municipal'])>0 else 0):.1f}% | Ejecutado: ${row['valor_calculado']:,.2f} | Estado: {row['estado']}"
                                )
                            desglose_operativo_txt = "\n".join(lineas_desglose)


                            
                            bloque_financiero_contexto = f"""
                            - Ámbito de Análisis: Entidad Territorial Local Unifamiliar [cite: 2089]
                            - Municipio Objeto de Auditoría: {muni_filtro} [cite: 2089]
                            - Total Asignado para el Municipio Específico: ${muni_total_asig:,.2f} [cite: 2090]
                            - Total Ejecutado y Aceptado para Pago del Municipio: ${muni_ejecutado_pago:,.2f} [cite: 2090]
                            - Total Pagado Efectivo Liquidado al Operador Municipal: ${muni_pagado_efectivo:,.2f} [cite: 2090, 2091]
                            - Saldo Contractual en Reserva Líquida Local: ${muni_saldo_reserva:,.2f} [cite: 2091]
                            """
                            
                            bloque_operativo_contexto = f"""
                            - Análisis Operativo Territorial: Eje de Metas de Subactividades Municipales
                            - Sumatoria Nominal de la Meta de la Subactividad Base (Teórica): {muni_meta_sub_nominal:,.1f} unidades.
                            - Sumatoria Nominal de Metas Municipales Programadas para este Contrato Local: {meta_programada:,.1f} unidades. [cite: 2092]
                            - Sumatoria de Avances Físicos Consolidados y Aprobados en Campo: {meta_avanzada:,.1f} unidades. [cite: 2093]
                            - Coeficiente de Eficiencia Física y Operativa Local: {porcentaje_operativo:.2f}% [cite: 2094]
                            - Matriz de Desglose y Trazabilidad Operativa (Actividad Maestro vs Subactividades Municipales):\n{desglose_operativo_txt} [cite: 2094]
                            """

                        # 3. COMPILACIÓN DE INDICADORES DE TRAZABILIDAD Y RESPUESTAS ADMINISTRATIVAS
                        cant_total_periodos = len(df_filtrado) [cite: 2094, 2095]
                        cant_plan_si = len(df_filtrado[df_filtrado['chk_plan_trabajo'] == 'Si']) if 'chk_plan_trabajo' in df_filtrado.columns else 0 [cite: 2095]
                        cant_crono_si = len(df_filtrado[df_filtrado['chk_cronograma'] == 'Si']) if 'chk_cronograma' in df_filtrado.columns else 0 [cite: 2095, 2096]
                        cant_personal_si = len(df_filtrado[df_filtrado['chk_personal'] == 'Si']) if 'chk_personal' in df_filtrado.columns else 0 [cite: 2096]
                        cant_segsoc_si = len(df_filtrado[df_filtrado['chk_seg_social'] == 'Si']) if 'chk_seg_social' in df_filtrado.columns else 0 [cite: 2096]
                        cant_iparcial_si = len(df_filtrado[df_filtrado['chk_inf_parcial'] == 'Si']) if 'chk_inf_parcial' in df_filtrado.columns else 0 [cite: 2096, 2097]
                        cant_ifinal_si = len(df_filtrado[df_filtrado['chk_inf_final'] == 'Si']) if 'chk_inf_final' in df_filtrado.columns else 0 [cite: 2097]
                        cant_polizas_si = len(df_filtrado[df_filtrado['chk_polizas'] == 'Si']) if 'chk_polizas' in df_filtrado.columns else 0 [cite: 2097]
                        
                        lista_obs_inf = df_filtrado['observaciones_referente'].dropna().astype(str).tolist() if 'observaciones_referente' in df_filtrado.columns else [] [cite: 2098]
                        bloque_observaciones_inf = "\n".join([f"- {o}" for o in lista_obs_inf if o.strip() != ""]) if lista_obs_inf else "Ninguna registrada." [cite: 2098]

                        # 4. CONSTRUCCIÓN DEL PROMPT DE CONTEXTO TRIPLE DIMENSIÓN CON BIFURCACIÓN PRECISA
                        contexto_ia = f"""
                        Actúa como un Consorcio Supervisor Experto de Alta Gerencia de Proyectos, Director Financiero de Contratación Estatal, Economista de la Salud y Científico de Datos. [cite: 2099]
                        Escribe el cuerpo analítico formal para el documento institucional: {nombre_informe_titulo}. [cite: 2100]
                        
                        1. ANÁLISIS ESTADÍSTICO-FINANCIERO Y COMPORTAMIENTO DE PAGOS:
                        {bloque_financiero_contexto} [cite: 2101]
                        - Ticket Promedio del Registro de Desembolsos en Periodo: ${promedio_pago:,.2f} [cite: 2101]
                        
                        2. CONVERGENCIA Y EVALUACIÓN OPERATIVA SANITARIA DE LAS METAS PIC:
                        {bloque_operativo_contexto}
                        
                        3. INFORME DE EVALUACIÓN Y CUMPLIMIENTO ADMINISTRATIVO (LISTA DE VERIFICACIÓN EN LA ACEPTACIÓN): [cite: 2102, 2103]
                        Analiza las respuestas de control diligenciadas obligatoriamente en cada radicación de cuenta:
                        - Entrega formal de Plan de Trabajo Concertado (Sí): {cant_plan_si} de {cant_total_periodos} registros. [cite: 2103]
                        - Entrega de Cronograma Mensualizado de campo (Sí): {cant_crono_si} de {cant_total_periodos} cortes. [cite: 2104]
                        - Verificación de idoneidad técnica del Personal contratado (Sí): {cant_personal_si} de {cant_total_periodos} líneas. [cite: 2105]
                        - Acreditación y pago de planillas de Seguridad Social (Sí): {cant_segsoc_si} de {cant_total_periodos} folios. [cite: 2106]
                        - Soporte de Informes Técnicos Parciales de Actividades (Sí): {cant_iparcial_si} de {cant_total_periodos} hitos. [cite: 2107]
                        - Entrega formal de Informes Técnicos Finales Consolidados (Sí): {cant_ifinal_si} de {cant_total_periodos} radicados. [cite: 2108]
                        - Vigencia de Pólizas de Cumplimiento Contractual (Sí): {cant_polizas_si} de {cant_total_periodos} amparos. [cite: 2109]
                        
                        ALERTAS Y NOVEDADES EXPUESTAS POR LOS REFERENTES TÉCNICOS:
                        {bloque_observaciones_inf} [cite: 2110]
                        
                        REQUISITOS DE REDACCIÓN COMPORTAMIENTO INSTITUCIONAL:
                        Escribe el informe de forma fluida, rigurosa, analítica y robusta, omitiendo resúmenes simples o listas planas. [cite: 2111] Utiliza terminología de auditoría fiscal, macroeconomía médica y epidemiología de control. [cite: 2112] Explica con total claridad la correlación que existe entre el cumplimiento de la lista de chequeo administrativa, la regularidad de giros y el avance real financiero/operativo del PIC. [cite: 2113]
                        Debe estructurar el desarrollo obligatorio en estas secciones:
                        - SECCIÓN I: DIAGNÓSTICO FINANCIERO Y ANÁLISIS LONGITUDINAL DE PAGOS. [cite: 2114]
                        - SECCIÓN II: CONVERGENCIA Y EVALUACIÓN OPERATIVA SANITARIA DE LAS METAS PIC.
                        - SECCIÓN III: AUDITORÍA Y CUMPLIMIENTO ADMINISTRATIVO LEGAL DEL CONTRATO. [cite: 2115]
                        - SECCIÓN IV: CONCLUSIONES GENERALES DE SUPERVISIÓN. [cite: 2116]
                        - SECCIÓN V: PLAN DE RECOMENDACIONES DE OPTIMIZACIÓN BASADAS EN EVIDENCIA (Mínimo 3 sugerencias). [cite: 2116]
                        """ [cite: 2117]






                        # 5. INTERVENCIÓN DIRECTA DEL ENDPOINT REST DE INTELIGENCIA ARTIFICIAL
                        try:
                            import requests
                            import json

                            api_key = st.secrets["GEMINI_API_KEY"]
                            url_api = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent?key={api_key}"
                            
                            headers = {"Content-Type": "application/json"}
                            payload = {
                                "contents": [{
                                    "parts": [{"text": contexto_ia}]
                                }]
                            }

                            res = requests.post(url_api, headers=headers, data=json.dumps(payload))
                            
                            if res.status_code == 200:
                                datos_json = res.json()
                                analisis_ia = datos_json['candidates'][0]['content']['parts'][0]['text']
                                st.session_state['ultimo_informe_tiempo'] = ahora
                            else:
                                raise Exception(f"Error de API externa (Código: {res.status_code})")
                                
                        # ✅ CÓDIGO CORREGIDO CON SANGRÍA EXACTA:
                        except Exception as e:
                            analisis_ia = f"""### EVALUACIÓN DE CONTINGENCIA: {nombre_informe_titulo}

El sistema PIC registra un índice de eficiencia financiera y operativa consolidada del {porcentaje_operativo:.2f}% bajo la cobertura analizada.
1. RESUMEN FINANCIERO: Los balances contables reportan un total asignado de ${(macro_total_asig if tipo_informe=='DEPARTAMENTAL' else muni_total_asig):,.2f}, logrando un presupuesto aceptado para cobro de ${macro_ejecutado_pago:,.2f} y giros efectivamente consolidados en banco de ${macro_pagado_efectivo:,.2f}.
El remanente financiero en reserva disponible se sitúa en ${(macro_reserva_disponible if tipo_informe=='DEPARTAMENTAL' else muni_saldo_reserva):,.2f}.
2. DESEMPEÑO ADMINISTRATIVO: Se auditaron {cant_total_periodos} hitos obligatorios de control administrativo, donde la tasa de acreditación de planillas de seguridad social y entrega formal de informes presenta un cumplimiento porcentual del {(cant_segsoc_si/cant_total_periodos*100) if cant_total_periodos > 0 else 0:.1f}%.
Se recomienda coordinar comités técnicos de supervisión inmediata para mitigar cuellos de botella contractuales."""

                        # 6. VISUALIZACIÓN EN PANTALLA (INTERFAZ DE USUARIO STREAMLIT)
                        st.success("📝 ¡Datos Procesados e Informe Estructurado con Éxito!")
                        
                        st.subheader("📊 Cuadro de Mando del Reporte Oficial")
                        inf_c1, inf_c2, inf_c3, inf_c4 = st.columns(4)
                        inf_c1.metric("Eficiencia Financiera", f"{porcentaje_ejecucion:.2f}%")
                        inf_c2.metric("Cumplimiento Operativo", f"{porcentaje_operativo:.2f}%")
                        inf_c3.metric("Ticket Promedio Ejecución", f"${promedio_pago:,.0f}")
                        inf_c4.metric("Registros Auditados", f"{cant_total_periodos} Hitos")

                        st.markdown("### 📈 Comparativo Estadístico de Rendimiento Real")
                        df_grafica_informe = pd.DataFrame({
                            "Dimensión Analizada": ["Rendimiento Financiero (Efectivo)", "Avance Operativo (Metas Físicas)"],
                            "Porcentaje de Cumplimiento (%)": [porcentaje_ejecucion, porcentaje_operativo]
                        })
                        st.bar_chart(data=df_grafica_informe, x="Dimensión Analizada", y="Porcentaje de Cumplimiento (%)")

                        st.markdown("### 🤖 Dictamen Analítico Avanzado Generado por IA")
                        st.markdown(analisis_ia)

                        # ==============================================================================
                        # 7. GENERACIÓN DEL DOCUMENTO WORD FORMAL (.DOCX) - INFORME TECNICO ADMINISTRATIVO Y FINANCIERO
                        # ==============================================================================
                        from docx import Document
                        from docx.shared import Inches, Pt
                        import io

                        doc = Document()
                        
                        for section in doc.sections:
                            section.top_margin = Inches(1)
                            section.bottom_margin = Inches(1)
                            section.left_margin = Inches(1)
                            section.right_margin = Inches(1)

                        title_p = doc.add_paragraph()
                        title_run = title_p.add_run("INFORME TECNICO ADMINISTRATIVO Y FINANCIERO")
                        title_run.bold = True
                        title_run.font.size = Pt(14)
                        title_run.font.name = 'Arial'
                        title_p.alignment = 1

                        doc.add_paragraph(f"Ámbito de Aplicación de Cobertura: {tipo_informe}")
                        if tipo_informe == "MUNICIPIO ESPECÍFICO":
                            doc.add_paragraph(f"Entidad Territorial Fiscalizada: {muni_filtro}")
                        doc.add_paragraph(f"Fecha de Emisión Automatizada: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')}")
                        doc.add_paragraph(f"Emitido por Profesional Supervisor: {st.session_state['user']}")
                        doc.add_paragraph("------------------------------------------------------------------------------------------------------------------------")

                        # TABLA DE INDICADORES FINANCIEROS DINÁMICOS
                        doc.add_heading('1. Matriz de Comportamiento Financiero y Flujo de Pagos', level=2)
                        table = doc.add_table(rows=1, cols=2)
                        table.style = 'Light Shading Accent 1'
                        hdr_cells = table.rows[0].cells
                        hdr_cells[0].text = 'Métrica Presupuestal y Económica Evaluada'
                        hdr_cells[1].text = 'Valor de Registro Oficial'

                        # 7.1 INYECCIÓN DE INDICADORES DE FILTRO EN TABLA WORD ASIMÉTRICA
                        if tipo_informe == "DEPARTAMENTAL":
                            métricas_word = [ [cite: 2143]
                                ("Techo Presupuestal Financiero Total Bolsa PIC", f"${macro_total_pic:,.2f}"), [cite: 2143]
                                ("Presupuesto Total Comprometido y Asignado a Municipios", f"${macro_total_asig:,.2f}"), [cite: 2143]
                                ("Presupuesto Ejecutado Aceptado para Pago", f"${macro_ejecutado_pago:,.2f}"), [cite: 2144]
                                ("Recursos Económicos Efectivamente Pagados (Giro Bancario OC)", f"${macro_pagado_efectivo:,.2f}"), [cite: 2144]
                                ("Saldo Libre de Reserva en Bolsa Central Departamental", f"${macro_reserva_disponible:,.2f}"), [cite: 2144]
                                ("Porcentaje de Municipios con Ejecución de Actividades", f"{tasa_cobertura_municipios:.1f}%"),
                                ("Sumatoria Techo de la 'Meta Global' Parametrizada", f"{meta_global_pic_nominal:,.0f} Unidades"),
                                ("Índice de Eficiencia Física en Meta Global PIC", f"{eficiencia_fisica_global:.2f}%")
                            ] [cite: 2145]
                        else:
                            métricas_word = [ [cite: 2146]
                                ("Presupuesto Total Estructural Asignado al Municipio", f"${muni_total_asig:,.2f}"), [cite: 2146]
                                ("Monto Local Ejecutado y Aceptado para Pago", f"${muni_ejecutado_pago:,.2f}"), [cite: 2146]
                                ("Monto Efectivamente Girado al Operador Local (Pagado OC)", f"${muni_pagado_efectivo:,.2f}"), [cite: 2147]
                                ("Saldo en Reserva Líquida Contractual Local Disponible", f"${muni_saldo_reserva:,.2f}"), [cite: 2147]
                                ("Sumatoria Meta Teórica de Subactividades Base", f"{muni_meta_sub_nominal:,.0f} Unidades"),
                                ("Metas Municipales Programadas para este Contrato Local", f"{meta_programada:,.0f} Unidades"), [cite: 2147]
                                ("Porcentaje de Eficiencia Física y Operativa Local", f"{porcentaje_operativo:.2f}%") [cite: 2148]
                            ]

                        else:
                            métricas_word = [
                                ("Presupuesto Total Estructural Asignado al Municipio", f"${muni_total_asig:,.2f}"),
                                ("Monto Local Ejecutado y Aceptado para Pago", f"${muni_ejecutado_pago:,.2f}"),
                                ("Monto Efectivamente Girado al Operador Local (Pagado OC)", f"${muni_pagado_efectivo:,.2f}"),
                                ("Saldo en Reserva Líquida Contractual Local Disponible", f"${muni_saldo_reserva:,.2f}"),
                                ("Porcentaje de Eficiencia Financiera Relativa", f"{porcentaje_ejecucion:.2f}%")
                            ]

                        for m, v in métricas_word:
                            row_cells = table.add_row().cells
                            row_cells[0].text = m
                            row_cells[1].text = v

                        doc.add_paragraph("\n")
                        
                        # TABLA DE SEGUIMIENTO ADMINISTRATIVO (RESPUESTAS DE ACEPTACIÓN)
                        doc.add_heading('2. Matriz de Cumplimiento Técnico, Trazabilidad y Control Administrativo', level=2)
                        table_adm = doc.add_table(rows=1, cols=3)
                        table_adm.style = 'Light Shading Accent 2'
                        hdr_adm_cells = table_adm.rows[0].cells
                        hdr_adm_cells[0].text = 'Criterio Administrativo Diligenciado en la Aceptación'
                        hdr_adm_cells[1].text = 'Casos Verificados (Sí)'
                        hdr_adm_cells[2].text = 'Efectividad (%)'
                        
                        items_admin_word = [
                            ("Radicación formal de Plan de Trabajo Concertado", f"{cant_plan_si} de {cant_total_periodos} registros", f"{(cant_plan_si / cant_total_periodos * 100) if cant_total_periodos > 0 else 0:.1f}%"),
                            ("Entrega y cumplimiento de Cronograma Mensualizado", f"{cant_crono_si} de {cant_total_periodos} registros", f"{(cant_crono_si / cant_total_periodos * 100) if cant_total_periodos > 0 else 0:.1f}%"),
                            ("Validación de idoneidad y perfiles del Personal", f"{cant_personal_si} de {cant_total_periodos} registros", f"{(cant_personal_si / cant_total_periodos * 100) if cant_total_periodos > 0 else 0:.1f}%"),
                            ("Verificación y pago de planillas de Seguridad Social", f"{cant_segsoc_si} de {cant_total_periodos} registros", f"{(cant_segsoc_si / cant_total_periodos * 100) if cant_total_periodos > 0 else 0:.1f}%"),
                            ("Radicación de Informes Técnicos Parciales de Soporte", f"{cant_iparcial_si} de {cant_total_periodos} registros", f"{(cant_iparcial_si / cant_total_periodos * 100) if cant_total_periodos > 0 else 0:.1f}%"),
                            ("Radicación formal de Informes Técnicos Finales", f"{cant_ifinal_si} de {cant_total_periodos} registros", f"{(cant_ifinal_si / cant_total_periodos * 100) if cant_total_periodos > 0 else 0:.1f}%"),
                            ("Vigencia de Amparos de Pólizas de Cumplimiento", f"{cant_polizas_si} de {cant_total_periodos} registros", f"{(cant_polizas_si / cant_total_periodos * 100) if cant_total_periodos > 0 else 0:.1f}%")
                        ]
                        
                        for crit, est, porc_t in items_admin_word:
                            row_adm = table_adm.add_row().cells
                            row_adm[0].text = crit
                            row_adm[1].text = est
                            row_adm[2].text = porc_t
                        
                        doc.add_paragraph("\n")
                        
                        # INYECCIÓN DEL ANÁLISIS ESTRATÉGICO INTEGRAL GENERADO CON LA IA
                        doc.add_heading('3. Diagnóstico de Situación, Conclusiones y Recomendaciones Basadas en Evidencia', level=2)
                        doc.add_paragraph(analisis_ia)

                        doc.add_paragraph("\n\n\n________________________________________________________")
                        doc.add_paragraph(f"Firma de Aval y Cierre Legal de Supervisión del Contrato\nPlan de Intervenciones Colectivas Departamental (Santander)")

                        bio = io.BytesIO()
                        doc.save(bio)
                        bio.seek(0)

                        st.download_button(
                            label="📥 Descargar INFORME TECNICO ADMINISTRATIVO Y FINANCIERO Oficial (.docx)",
                            data=bio,
                            file_name=f"INFORME_TECNICO_ADMINISTRATIVO_Y_FINANCIERO_PIC_{tipo_informe}_{pd.Timestamp.now().strftime('%Y%m%d')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="btn_download_word_v66_oficial"
                        )
# ==============================================================================


# --- CONSOLIDADO GLOBAL DE PAGOS (VISTA DEPARTAMENTO) ---
        st.write("---")
        st.subheader("📑 Trazabilidad General de Pagos")
        
        # Consolidado global uniendo hojas de Excel
        df_p_glob = get_data("seguimiento_pagos")
        df_a_glob = get_data("asignacion_municipios")
        df_s_glob = get_data("subactividades")
        
        if not df_p_glob.empty and not df_a_glob.empty and not df_s_glob.empty and not df_act_raw.empty:
            # Homologación defensiva del nombre de la variable monetaria
            df_p_copy = df_p_glob.copy()
            if 'valor_calculated' in df_p_copy.columns and 'valor_calculado' not in df_p_copy.columns:
                df_p_copy.rename(columns={'valor_calculated': 'valor_calculado'}, inplace=True)

            # Re-ensamble matricial completo usando "left joins" para proteger la integridad de filas
            df_global = df_p_copy.merge(df_a_glob, on="id_asig", how="left")
            df_global = df_global.merge(df_s_glob, on="id_sub", how="left")
            df_global = df_global.merge(df_act_raw, on="id_actividad", how="left")

            
            # Reorganización explícita de columnas inyectando los datos de la Actividad Padre
            cols_verificar = [
                'id_seguimiento', 'municipio', 'id_actividad', 'nombre_actividad', 
                'nombre_subactividad', 'num_pago_actual', 'valor_calculado', 
                'referente_aprobador', 'acta_referente', 'supervisor_aprobador', 'estado', 'motivo_rechazo'
            ]
            
            for col in cols_verificar:
                if col not in df_global.columns:
                    df_global[col] = None

            # Filtrar el DataFrame final con el orden estricto establecido
            df_global = df_global[cols_verificar].copy()
            
            # Tratamiento y limpieza de registros vacíos
            df_global['id_actividad'] = df_global['id_actividad'].fillna("N/A")
            df_global['nombre_actividad'] = df_global['nombre_actividad'].fillna("Sin Actividad Padre")
            df_global['referente_aprobador'] = df_global['referente_aprobador'].fillna("Pendiente Aval")
            df_global['acta_referente'] = df_global['acta_referente'].fillna("Sin Acta")
            df_global['supervisor_aprobador'] = df_global['supervisor_aprobador'].fillna("Pendiente Firma")
            df_global['motivo_rechazo'] = df_global['motivo_rechazo'].fillna("Sin observaciones")
            df_global['estado'] = df_global['estado'].fillna("PENDIENTE")
            
            # Asignación final de encabezados profesionales e independientes para Streamlit
            df_global.columns = [
                'ID Seguimiento', 'Municipio', 'ID Actividad', 'Actividad', 
                'Sub Actividad', 'Pago N°', 'Valor Proporcional', 'Referente que Avaló', 
                'Acta Referente', 'Supervisor que Aprobó', 'Estado', 'Observaciones / Motivo de Rechazo'
            ]
            
            # Formateador de moneda dinámico y seguro
            df_global['Valor Proporcional'] = pd.to_numeric(df_global['Valor Proporcional'], errors='coerce').fillna(0)
            df_global['Valor Proporcional'] = df_global['Valor Proporcional'].map(lambda x: f"${x:,.2f}")
        else:
            df_global = pd.DataFrame()

        if not df_global.empty:
            st.dataframe(df_global, use_container_width=True)

# --- MÓDULO: GESTIÓN DE USUARIOS ---
    elif menu == "👤 Gestión Usuarios":
        if rol != "DEPARTAMENTO_PARAMETRIZADOR":
            st.warning("Acceso denegado.")
        else:
            st.title("👤 Administración de Usuarios")
            with st.form("crear_usuario"):
                c1, c2 = st.columns(2)
                u_nombre = c1.text_input("Nombre Completo")
                u_email = c1.text_input("Correo Electrónico (Usuario)")
                u_pass = c1.text_input("Contraseña", type="password")
                u_cedula = c2.text_input("Cédula")
                u_cargo = c2.text_input("Cargo")
                u_tel = c2.text_input("Teléfono")
                u_rol = st.selectbox("Asignar Rol", ["DEPARTAMENTO_PARAMETRIZADOR", "MUNICIPIO_EJECUTOR", "REFERENTE_DEPARTAMENTAL", "SUPERVISOR"])
                u_muni = st.selectbox("Municipio Asignado", ["N/A"] + municipios_santander)

                if st.form_submit_button("Registrar Usuario"):
                    df_users = get_data("usuarios")
                    if u_email in df_users['email'].values:
                        st.error("Error: El correo ya existe.")
                    else:
                        nueva_fila = pd.DataFrame([{
                            "id_usuario": len(df_users) + 1,
                            "email": u_email, "password": u_pass, "nombre_completo": u_nombre,
                            "cedula": u_cedula, "cargo": u_cargo, "telefono": u_tel,
                            "rol": u_rol, "municipio_asignado": u_muni
                        }])
                        df_final = pd.concat([df_users, nueva_fila], ignore_index=True)
                        safe_update("usuarios", df_final)
                        st.success(f"Usuario {u_email} registrado correctamente.")
                        st.rerun()

# --- VISUALIZACIÓN Y GESTIÓN DE USUARIOS EXISTENTES ---
            st.write("---")
            st.subheader("👥 Usuarios Registrados")
            
            
            df_usuarios = get_data("usuarios")
            
            if not df_usuarios.empty:
                st.dataframe(df_usuarios, use_container_width=True)
                
                col_edit, col_del = st.columns(2)
                
                # --- OPCIÓN: ACTUALIZAR ---
                with col_edit.expander("📝 Editar Usuario"):
                    id_update = st.number_input("ID del usuario a editar:", min_value=1, step=1, key="upd_user_id")
                    user_to_edit = df_usuarios[df_usuarios['id_usuario'] == id_update]
                    
                    if not user_to_edit.empty:
                        new_nombre = st.text_input("Nuevo Nombre", value=user_to_edit.iloc[0]['nombre_completo'])
                        new_rol = st.selectbox("Nuevo Rol", ["DEPARTAMENTO_PARAMETRIZADOR", "MUNICIPIO_EJECUTOR", "REFERENTE_DEPARTAMENTAL", "SUPERVISOR"], 
                                             index=["DEPARTAMENTO_PARAMETRIZADOR", "MUNICIPIO_EJECUTOR", "REFERENTE_DEPARTAMENTAL", "SUPERVISOR"].index(user_to_edit.iloc[0]['rol']))
                        
                        if st.button("Guardar Cambios"):
                            df_u = get_data("usuarios")
                            df_u.loc[df_u['id_usuario'] == id_update, ['nombre_completo', 'rol']] = [new_nombre, new_rol]
                            safe_update("usuarios", df_u)
                            st.success("✅ Usuario actualizado en Excel.")
                            st.rerun()

                # --- OPCIÓN: ELIMINAR ---
                with col_del.expander("🗑️ Eliminar Usuario"):
                    id_delete = st.number_input("ID del usuario a eliminar:", min_value=1, step=1, key="del_user_id")
                    if st.button("Confirmar Eliminación de Usuario", type="primary"):
                        if id_delete == 1: # Protección para no borrar al admin principal
                            st.error("No se puede eliminar el usuario administrador maestro.")
                        else:
                            df_u_del = get_data("usuarios")
                            df_u_del = df_u_del[df_u_del['id_usuario'] != id_delete]
                            safe_update("usuarios", df_u_del)
                            st.warning(f"⚠️ Usuario {id_delete} eliminado del Excel")
                            st.rerun()
            else:
                st.info("No hay usuarios registrados además del administrador.")
